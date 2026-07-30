"""Populated DOCX: weekly highlight report (EN), IMO Wochenprotokoll (DE),
WS1 Teilprojekt-Protokoll (DE), Expertenrunde ERP (DE), Rollenkarten (DE)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import case as C

STATUS_DE = {"Open": "Offen", "In progress": "In Arbeit", "Done": "Erledigt",
             "Blocked": "Blockiert", "Overdue": "Überfällig"}
MS_EN = {"done": "Done", "on_track": "On track", "at_risk": "At risk", "delayed": "Delayed"}
MS_DE = {"done": "Erledigt", "on_track": "Im Plan", "at_risk": "Gefährdet",
         "delayed": "Verzögert"}


def doc():
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(9.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for s in d.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
    for lvl, sz in ((1, 14), (2, 11.5), (3, 10)):
        h = d.styles[f"Heading {lvl}"]
        h.font.name = "Arial"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.color.rgb = RGBColor.from_string(C.D_DARK if lvl < 3 else C.D_BLACK)
    return d


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexcolor)
    tcPr.append(el)


def kv(d, pairs):
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in pairs:
        row = t.add_row()
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(12.0)
        shade(row.cells[0], C.D_PALE)
        r = row.cells[0].paragraphs[0].add_run(str(k))
        r.bold = True
        r.font.size = Pt(8.5)
        r2 = row.cells[1].paragraphs[0].add_run(str(v))
        r2.font.size = Pt(8.5)
    d.add_paragraph()
    return t


def grid(d, headers, rows, widths=None, red_cols=()):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, C.D_DARK)
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        if widths:
            c.width = Cm(widths[j])
    for row in rows:
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(7.5)
            if j in red_cols and str(val) in ("Delayed", "Verzögert", "Überfällig", "Overdue",
                                             "Blocked", "Blockiert", "Red", "Rot"):
                r.font.color.rgb = RGBColor.from_string(C.RAG_COLOR["Red"])
                r.bold = True
            elif j in red_cols and str(val) in ("At risk", "Gefährdet", "Amber", "Gelb"):
                r.font.color.rgb = RGBColor.from_string(C.RAG_COLOR["Amber"])
                r.bold = True
            if widths:
                c.width = Cm(widths[j])
    d.add_paragraph()
    return t


def head(d, title, subtitle, pairs):
    p = d.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(C.D_DARK)
    p2 = d.add_paragraph()
    r = p2.add_run(subtitle)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(C.D_GREY)
    d.add_paragraph()
    kv(d, pairs)


def note(d, text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(C.D_GREY)


def para(d, text, size=9.5):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def disclaimer(d, german=False):
    d.add_paragraph()
    if german:
        note(d, "Synthetisches Dokument. Der Fallbezug (Unternehmen, Transaktionswert, "
                "Abschlussdatum, Governance-Struktur, Teilprojektzuschnitt) folgt oeffentlich "
                "verfuegbaren Quellen. Alle operativen Namen, Zahlen, Risiken, Beschluesse und "
                "Synergiewerte sind erfunden, in sich konsistent und stellen keine Aussage "
                "darueber dar, was die beteiligten Parteien tatsaechlich getan haben.")
    else:
        note(d, "Synthetic document. The case anchoring (entities, transaction value, close "
                "date, governance structure, workstream cut) follows publicly available "
                "sources. All operational names, figures, risks, decisions and synergy values "
                "are invented, internally consistent, and are not a representation of what any "
                "party actually did.")


# =====================================================================
# 1. Weekly Integration Highlight Report (EN)
# =====================================================================
def highlight():
    d = doc()
    head(d, f"{C.OFFICE} - Weekly Integration Highlight Report",
         f"{C.ACQUIRER} and {C.TARGET}  |  Reporting week {C.WEEK_LABEL}  |  "
         f"Day {C.DAYS_AFTER_DAY1} after Day 1",
         [("Document", "Weekly Integration Highlight Report"),
          ("Programme", f"{C.PROGRAM}, run by the {C.OFFICE} ({C.OFFICE_ABBR})"),
          ("Reporting period", f"{C.en(C.WEEK_START)} to {C.en(C.WEEK_END)}"),
          ("Prepared by", f"{C.nm('imo_pmo')}, {C.role('imo_pmo')}"),
          ("Reviewed by", f"{C.nm('prog_dir')}, {C.role('prog_dir')}"),
          ("Distribution", "Steering Committee, workstream leads, executive sponsor"),
          ("Classification", "Strictly confidential"),
          ("Version and status", "v1.0, issued")])

    d.add_heading("1. Overall assessment", level=1)
    reds = [C.ws_full(c) for c in C.WS_CODES if C.ws_rag(c) == "Red"]
    para(d, f"The programme is Amber. The Day 100 commitment of {C.en(C.DAY100)} is still "
            f"credible, but for the first time since close three workstreams carry a delayed "
            f"gate milestone at the same time: {', '.join(reds)}. Two of those delays share a "
            f"single root cause, which is that decisions taken outside the integration "
            f"programme, namely the legacy ERP upgrade at the target and the pace of the EMEA "
            f"entity review, are setting the pace of work inside it.")
    para(d, f"Nothing this week requires the Day 100 date to move. What it does require is a "
            f"decision on scope, which is why B-07 is on the Steering Committee agenda for "
            f"{C.en(C.STEERCO_02)}.")
    kv(d, [("Overall RAG this week", C.OVERALL_RAG),
           ("Overall RAG last week", "Green"),
           ("Reason for the change", f"M-05 and M-07 both moved from at risk to delayed, and "
                                     f"M-13 slipped {(C.MILESTONES[12][7] - C.MILESTONES[12][6]).days} days"),
           ("Milestones on track", f"{C.MS_ON_TRACK} of {C.MS_TOTAL} "
                                   f"({C.MS_DONE} complete, {C.MS_DELAYED} delayed, "
                                   f"{C.MS_AT_RISK} at risk)"),
           ("Tasks", f"{C.T_TOTAL} tracked, {C.T_DONE} done, {C.T_INPROG} in progress, "
                     f"{C.T_OPEN} open, {C.T_OVERDUE} overdue, {C.T_BLOCKED} blocked"),
           ("Overall progress", f"{C.OVERALL_PROGRESS} percent, weighted mean across all tasks"),
           ("Open risks", f"{C.R_TOTAL}, of which {C.R_HIGH} at or above the escalation "
                          f"threshold of severity {C.ESCALATION_THRESHOLD}"),
           ("Synergy secured", f"USD {C.SYN_SECURED} m of USD {C.SYN_TARGET} m register target "
                               f"({C.SYN_SECURED_PCT} percent), of which USD {C.SYN_VALIDATED} m "
                               f"is Finance validated")])

    d.add_heading("2. Achievements this week", level=1)
    for t in [t for t in C.TASKS if t[6] == "Done"]:
        d.add_paragraph(f"{t[3]} {C.WS_NAME[t[3]]}: {t[1]} (task {t[0]}, "
                        f"owner {C.nm(t[4])}, closed by {C.en(t[5])})", style="List Bullet")
    d.add_paragraph(f"Steering Committee decision B-06 took effect: from session 02 onward only "
                    f"Finance-validated initiatives count towards secured synergy, which closed "
                    f"issue I-03.", style="List Bullet")

    d.add_heading("3. Status by workstream", level=1)
    rows = []
    for code in C.WS_CODES:
        ms = C.ws_milestones(code)
        done = [t for t in C.ws_tasks(code) if t[6] == "Done"]
        nxt = [t for t in C.ws_tasks(code) if t[6] in ("Open", "In progress")][:1]
        rows.append([code, C.WS_NAME[code], C.nm(code), C.ws_rag(code),
                     f"{C.ws_progress(code)}%",
                     done[0][0] if done else "none this week",
                     nxt[0][0] if nxt else "-",
                     f"{sum(1 for m in ms if m[8] == 'delayed')} delayed of {len(ms)}"])
    grid(d, ["WS", "Workstream", "Lead", "RAG", "Progress", "Closed this week",
             "Next due", "Milestones"], rows,
         [1.2, 2.8, 2.4, 1.4, 1.6, 2.4, 1.6, 3.2], red_cols=(3,))

    d.add_heading("4. Deviations and delays", level=1)
    rows = []
    for m in C.MILESTONES:
        if m[7] != m[6]:
            rows.append([m[0], m[1][:56], m[3], C.nm(m[4]) if m[4] in C.PEOPLE else "",
                         C.en(m[6]), C.en(m[7]), (m[7] - m[6]).days, m[10][:70]])
    grid(d, ["ID", "Milestone", "WS", "Owner", "Baseline", "Forecast", "Slip (days)",
             "Root cause and recovery"], rows,
         [1.2, 4.0, 1.0, 2.2, 2.0, 2.0, 1.4, 3.8])

    d.add_heading("5. Risks at or above the escalation threshold", level=1)
    rows = []
    for r in sorted([r for r in C.RISKS if C.sev(r) >= C.ESCALATION_THRESHOLD],
                    key=C.sev, reverse=True):
        rows.append([r[0], r[1][:70], r[3], f"{r[5]} x {r[6]} = {C.sev(r)}", C.band(C.sev(r)),
                     r[10], r[7][:66], C.nm(r[4]), C.en(r[8])])
    grid(d, ["ID", "Risk", "WS", "L x I", "Band", "Trend", "Mitigation", "Owner", "Due"],
         rows, [1.0, 3.6, 0.9, 1.4, 1.1, 1.6, 3.4, 2.0, 1.8])
    note(d, f"Escalation rule from the integration charter: severity {C.ESCALATION_THRESHOLD} "
            f"and above goes to the Steering Committee. {C.R_ESCALATED} risks currently sit "
            f"above the workstream level.")

    d.add_heading("6. Decisions taken", level=1)
    rows = [[x[0], x[1][:74], x[3][:66], x[4], C.en(x[5]), x[6], x[8]]
            for x in C.DECISIONS if x[8] != "Open"]
    grid(d, ["ID", "Decision", "Rationale", "Body", "Date", "Affects", "Status"],
         rows, [1.0, 4.4, 4.0, 1.8, 2.0, 2.0, 1.8])

    d.add_heading("7. Escalations to the Steering Committee", level=1)
    rows = [[x[0], x[1][:78], x[3][:70], C.nm(x[7]), C.en(C.STEERCO_02)]
            for x in C.DECISIONS if x[8] == "Open"]
    grid(d, ["ID", "Requested decision", "Why it needs the committee", "Proposed by",
             "Needed by"], rows, [1.0, 5.0, 4.6, 2.4, 2.0])

    d.add_heading("8. Look ahead: next two weeks", level=1)
    rows = []
    for m in sorted(C.MILESTONES, key=lambda x: x[7]):
        if C.TODAY < m[7] <= C.DELL_EMC_WORLD:
            rows.append([m[0], m[1][:60], m[3], C.nm(m[4]) if m[4] in C.PEOPLE else "",
                         C.en(m[7]), MS_EN[m[8]], "Yes" if m[9] else "No"])
    grid(d, ["ID", "Milestone", "WS", "Owner", "Due", "Readiness", "Gate"], rows,
         [1.0, 4.6, 1.0, 2.4, 2.2, 1.8, 1.2], red_cols=(5,))
    para(d, f"The flagship customer event on {C.en(C.DELL_EMC_WORLD)} is an external fixed date. "
            f"Everything in the table above that carries a gate flag has to be true before it.")

    d.add_heading("9. Annexes", level=1)
    for a in [f"Annex A: integration tracker as at {C.en(C.D_TRACKER)}, xlsx",
              f"Annex B: RAID log as at {C.en(C.D_TRACKER)}, xlsx",
              f"Annex C: synergy tracker as at {C.en(C.D_SYNERGY)}, xlsx",
              f"Annex D: Steering Committee pack session 02, {C.en(C.STEERCO_02)}, pptx"]:
        d.add_paragraph(a, style="List Bullet")

    disclaimer(d)
    d.save(C.OUT / "DellEMC_VCIO_Weekly_Highlight_Report_W3_2016-09-30.docx")


# =====================================================================
# 2. IMO Wochenprotokoll (DE)
# =====================================================================
def protokoll_imo():
    d = doc()
    head(d, f"{C.OFFICE} ({C.OFFICE_ABBR}) - Protokoll der Wochenbesprechung",
         f"{C.ACQUIRER} und {C.TARGET}  |  {C.CW} 2016  |  "
         f"Tag {C.DAYS_AFTER_DAY1} nach Day 1  |  Regionale Sicht DACH",
         [("Dokument", "Protokoll der woechentlichen Integrationsbesprechung"),
          ("Programm", f"{C.PROGRAM}"),
          ("Sitzung", "VCIO Weekly, regionale Sitzung DACH"),
          ("Datum und Uhrzeit", f"{C.de(C.D_PROTOKOLL)}, 09:00 bis 10:30 Uhr MESZ"),
          ("Ort", "Videokonferenz, Bruecke Frankfurt und Schwalbach"),
          ("Sitzungsleitung", f"{C.nm('dach_lead')}, {C.role('dach_lead')}"),
          ("Protokollfuehrung", f"{C.nm('imo_pmo')}, {C.role('imo_pmo')}"),
          ("Verteiler", "Alle Teilprojektleitungen, VCIO Programmleitung, "
                        "Regionalleitung DACH"),
          ("Klassifizierung", "Streng vertraulich"),
          ("Version und Status", "v1.0, freigegeben")])

    d.add_heading("1. Teilnehmer", level=1)
    att = [(C.nm(k), C.role(k), ws, "ja") for k, ws in
           [("dach_lead", "-"), ("imo_pmo", "-"), ("imo_mgr", "-"), ("WS1", "WS1"),
            ("WS2", "WS2"), ("WS3", "WS3"), ("dach_hr", "WS4"), ("WS5", "WS5"),
            ("WS6", "WS6"), ("advisor", "-")]]
    att.append((C.nm("WS4"), C.role("WS4"), "WS4", "nein, vertreten durch " + C.nm("dach_hr")))
    att.append((C.nm("WS7"), C.role("WS7"), "WS7", "nein, entschuldigt"))
    grid(d, ["Name", "Funktion", "Teilprojekt", "Anwesend"], att,
         [3.4, 6.0, 2.4, 5.2])

    d.add_heading("2. Agenda", level=1)
    grid(d, ["Nr.", "Agendapunkt", "Verantwortlich", "Minuten", "Art"], [
        ["1", "Statusrunde der Teilprojekte", C.nm("imo_pmo"), "30", "Information"],
        ["2", "Meilenstein M-07, Zielbild ERP-Konsolidierung", C.nm("WS3"), "20", "Diskussion"],
        ["3", "Meilenstein M-13, Anhoerung der Betriebsraete DACH", C.nm("br_liaison"),
         "20", "Beschluss"],
        ["4", "Vorbereitung der Steering-Committee-Sitzung 02", C.nm("imo_mgr"),
         "15", "Diskussion"],
        ["5", "Offene Punkte aus der Vorwoche", C.nm("imo_pmo"), "5", "Information"],
    ], [1.2, 7.0, 3.4, 1.8, 3.6])

    d.add_heading("3. Statusrunde Teilprojekte", level=1)
    note(d, f"Kurzstatus je Teilprojekt. Fuehrende Quelle fuer Aufgaben und Termine ist der "
            f"Integration Tracker, Stand {C.de(C.D_TRACKER)}, gemaess Beschluss B-02.")
    rows = []
    for code in C.WS_CODES:
        rag = {"Green": "Gruen", "Amber": "Gelb", "Red": "Rot"}[C.ws_rag(code)]
        ms = C.ws_milestones(code)
        delayed = [m[0] for m in ms if m[8] == "delayed"]
        rows.append([code, C.WS_NAME_DE[code], C.nm(code), rag, f"{C.ws_progress(code)} %",
                     str(C.ws_overdue(code)),
                     ", ".join(delayed) if delayed else "keiner"])
    grid(d, ["Nr.", "Teilprojekt", "Leitung", "Ampel", "Fortschritt", "Ueberfaellig",
             "Verzoegerte Meilensteine"], rows,
         [1.2, 3.4, 2.8, 1.4, 2.0, 2.2, 4.0], red_cols=(3,))

    d.add_heading("4. Diskussion und Ergebnisse", level=1)

    d.add_heading("4.1 Meilenstein M-07, Zielbild ERP-Konsolidierung", level=2)
    para(d, f"{C.nm('WS3')} berichtet, dass das Zielbild der ERP-Konsolidierung nicht zum "
            f"Basistermin {C.de(C.MILESTONES[6][6])} freigegeben werden kann. Ursache ist "
            f"nicht die Integrationsarbeit selbst, sondern die parallel laufende Modernisierung "
            f"des Altsystems beim Zielunternehmen: beide Vorhaben benoetigen dieselben "
            f"Architekturentscheidungen und dieselben Personen.")
    para(d, f"Als Prognosetermin wird {C.de(C.M07_MINUTES)} festgehalten. "
            f"{C.nm('dach_it')} weist darauf hin, dass dieser Termin nur dann haltbar ist, wenn "
            f"die erneute Extraktion der Materialstammdaten bis {C.de(C.MILESTONES[6][7])} "
            f"verwertbar vorliegt; andernfalls verschiebt sich das Zielbild ein zweites Mal.")
    para(d, f"{C.nm('WS5')} hebt hervor, dass die Abhaengigkeit D-02 damit unmittelbar betroffen "
            f"ist. Die Lieferkette kann die Stammdatenbereinigung nicht abschliessen, solange "
            f"das Zielsystem nicht entschieden ist.")
    note(d, "Anmerkung der Protokollfuehrung: Der hier festgehaltene Prognosetermin wurde nach "
            "der Sitzung im Mailverlauf erneut angepasst. Das Protokoll wurde nicht nachgezogen.")

    d.add_heading("4.2 Meilenstein M-13, Anhoerung der Betriebsraete DACH", level=2)
    para(d, f"{C.nm('br_liaison')} berichtet, dass die Anhoerung nach den einschlaegigen "
            f"Vorschriften des Betriebsverfassungsgesetzes nicht bis zum Basistermin "
            f"{C.de(C.MILESTONES[12][6])} abgeschlossen werden kann. Der Prognosetermin lautet "
            f"{C.de(C.MILESTONES[12][7])}.")
    para(d, f"{C.nm('dach_hr')} stellt klar, dass es sich um ein rechtliches Tor handelt und "
            f"nicht um eine Kommunikationspraeferenz. Eine Veroeffentlichung der Zielstruktur "
            f"vor Abschluss der Anhoerung wuerde ein Compliance-Risiko erzeugen.")
    para(d, f"{C.nm('WS6')} weist auf die Folgewirkung hin: das Verguetungsmodell im Vertrieb "
            f"(M-09) haengt an derselben Freigabe, und davon wiederum die Gebiets- und "
            f"Quotenzuordnung.")

    d.add_heading("4.3 Vorbereitung der Steering-Committee-Sitzung 02", level=2)
    para(d, f"{C.nm('imo_mgr')} fasst zusammen: dem Steering Committee wird am "
            f"{C.de(C.STEERCO_02)} ein Beschluss vorgelegt, den Termin des "
            f"Wertrealisierungsberichts zu Day 100 zu halten und stattdessen den Umfang neu zu "
            f"takten. Die Alternative, den Termin zu verschieben, wird ausdruecklich nicht "
            f"empfohlen, weil es sich um eine Zusage an den Sponsor handelt.")
    para(d, f"{C.nm('advisor')} ergaenzt, dass die Priorisierung nach dem Wertpriorisierungs"
            f"rahmen bereits die Grundlage liefert, um zu bestimmen, welcher Umfang verschoben "
            f"werden kann, ohne den wesentlichen Teil des Wertbeitrags zu gefaehrden.")

    d.add_heading("5. Beschluesse", level=1)
    rows = [[x[0], x[2][:80], x[3][:70], x[4], C.de(x[5]), x[6]]
            for x in C.DECISIONS if x[5] <= C.D_PROTOKOLL]
    grid(d, ["Nr.", "Beschluss", "Begruendung", "Gremium", "Datum", "Betroffen"],
         rows, [1.2, 5.0, 4.4, 2.0, 2.0, 2.4])

    d.add_heading("6. Aufgaben und Massnahmen", level=1)
    rows = [[a[0], a[2][:74], C.nm(a[4]), a[5], C.de(a[7]), STATUS_DE[a[9]]]
            for a in C.ACTIONS]
    grid(d, ["Nr.", "Aufgabe", "Verantwortlich", "Teilprojekt", "Faellig am", "Status"],
         rows, [1.4, 6.4, 2.8, 2.0, 2.2, 2.2], red_cols=(5,))

    d.add_heading("7. Risiken und Eskalationen", level=1)
    rows = []
    for r in sorted(C.RISKS, key=C.sev, reverse=True)[:6]:
        rows.append([r[0], r[2][:74], f"{r[5]}", f"{r[6]}", f"{C.sev(r)}",
                     C.nm(r[4]), r[11] if r[11] else "keine"])
    grid(d, ["Nr.", "Risiko", "Wahrsch.", "Auswirkung", "Schwere", "Verantwortlich",
             "Eskalation an"], rows, [1.2, 6.0, 1.6, 1.8, 1.6, 2.6, 2.2])
    note(d, f"Eskalationsschwelle laut Integration Charter: ab Schwere "
            f"{C.ESCALATION_THRESHOLD} an das Steering Committee. Derzeit erfuellen "
            f"{C.R_HIGH} Risiken diese Schwelle.")

    d.add_heading("8. Offene Punkte aus der Vorwoche", level=1)
    rows = [[a[0], a[2][:70], C.nm(a[4]), C.de(a[6]), C.de(a[7]), str(a[8]), STATUS_DE[a[9]]]
            for a in C.ACTIONS if a[8] > 0]
    grid(d, ["Nr.", "Offener Punkt", "Verantwortlich", "Urspruenglich faellig",
             "Neuer Termin", "Verschiebungen", "Status"], rows,
         [1.2, 5.2, 2.6, 2.6, 2.2, 2.0, 2.2])

    d.add_heading("9. Naechste Sitzung", level=1)
    kv(d, [("Termin", "04.10.2016, 09:00 Uhr MESZ"),
           ("Schwerpunkt", "Stand M-07 nach erneuter Stammdatenextraktion, "
                           "Vorbereitung Day 30"),
           ("Vorzubereiten von", f"{C.nm('WS3')} und {C.nm('dach_it')}")])

    disclaimer(d, german=True)
    d.save(C.OUT / "DellEMC_VCIO_Wochenprotokoll_KW39_2016-09-27.docx")


# =====================================================================
# 3. Teilprojekt-Protokoll Finance (DE)
# =====================================================================
def protokoll_ws1():
    d = doc()
    head(d, f"{C.PROGRAM} - Protokoll Teilprojekt WS1 Finanzen",
         f"Jour Fixe {C.CW} 2016  |  Regionale Sicht DACH  |  "
         f"Tag {C.DAYS_AFTER_DAY1} nach Day 1",
         [("Dokument", "Protokoll der Teilprojektbesprechung WS1 Finanzen"),
          ("Programm", C.PROGRAM),
          ("Teilprojektleitung", f"{C.nm('WS1')}, {C.role('WS1')}"),
          ("Datum und Uhrzeit", f"{C.de(C.D_PROTOKOLL)}, 14:00 bis 15:30 Uhr MESZ"),
          ("Ort", "Videokonferenz"),
          ("Protokollfuehrung", f"{C.nm('dach_fin')}, {C.role('dach_fin')}"),
          ("Verteiler", "Teilprojektteam WS1, VCIO, Konzernrechnungswesen"),
          ("Klassifizierung", "Streng vertraulich"),
          ("Version und Status", "v1.0, freigegeben")])

    d.add_heading("1. Teilnehmer", level=1)
    grid(d, ["Name", "Rolle", "Gesellschaft", "Anwesend"], [
        [C.nm("WS1"), C.role("WS1"), C.NEWCO, "ja"],
        [C.nm("dach_fin"), C.role("dach_fin"), "Dell EMC DACH", "ja"],
        [C.nm("synergy"), C.role("synergy"), C.NEWCO, "ja"],
        [C.nm("WS2"), C.role("WS2"), C.NEWCO, "ja, zu Punkt 3.2"],
        [C.nm("advisor"), C.role("advisor"), C.ADVISOR, "ja"],
    ], [3.4, 5.6, 3.6, 4.4])

    d.add_heading("2. Statusuebersicht Arbeitspakete", level=1)
    rows = [[t[0], t[2][:66], C.nm(t[4]), C.de(t[5]), STATUS_DE[t[6]], f"{t[7]} %",
             t[9] if t[9] else "keine"] for t in C.ws_tasks("WS1")]
    grid(d, ["Nr.", "Arbeitspaket", "Verantwortlich", "Faellig am", "Status", "Fortschritt",
             "Abhaengig von"], rows, [1.2, 5.4, 2.6, 2.0, 2.0, 1.8, 2.0], red_cols=(4,))
    note(d, f"Teilprojektfortschritt gesamt {C.ws_progress('WS1')} Prozent, "
            f"gewichteter Mittelwert ueber {len(C.ws_tasks('WS1'))} Arbeitspakete. "
            f"Ampel {('Gruen' if C.ws_rag('WS1') == 'Green' else 'Gelb')}.")

    d.add_heading("3. Fachliche Themen", level=1)

    d.add_heading("3.1 Kontenplan und Abschlusskalender", level=2)
    para(d, f"Die Abbildung des Kontenplans (A-001) liegt bei 70 Prozent. Die Angleichung des "
            f"Abschlusskalenders (A-002) ist mit 85 Prozent weiter, weil sie keine Abhaengigkeit "
            f"nach aussen hat. {C.nm('dach_fin')} weist darauf hin, dass die Abstimmung der "
            f"konzerninternen Salden (A-005) erst beginnen kann, wenn die Kontenzuordnung steht.")

    d.add_heading("3.2 Konsolidierungskreis und Gesellschaftsstruktur", level=2)
    para(d, f"{C.nm('WS2')} berichtet, dass die steuerliche Pruefung der EMEA-Gesellschaftskette "
            f"(A-011) ihren Termin ueberschritten hat und damit den Freigabeplan M-05 blockiert. "
            f"Solange die Struktur nicht freigegeben ist, kann der Konsolidierungskreis fuer das "
            f"erste gemeinsame Quartal (A-003) nicht festgelegt werden. Das ist die Abhaengigkeit "
            f"D-01, die bislang nur von einer Seite bestaetigt ist.")
    para(d, f"Vereinbart wird, die Pruefung an das Steering Committee zu eskalieren und einen "
            f"verbindlichen Abschlusstermin zu benennen (OP-05).")

    d.add_heading("3.3 Synergie-Baseline und Validierung", level=2)
    para(d, f"{C.nm('synergy')} berichtet den Stand der Validierung: von USD "
            f"{C.SYN_SECURED} Mio. gesicherter Synergie sind USD {C.SYN_VALIDATED} Mio. durch "
            f"Finance validiert. Der Unterschied von USD {C.SYN_SECURED - C.SYN_VALIDATED} Mio. "
            f"entfaellt vollstaendig auf Umsatzsynergien, fuer die bislang keine Initiative "
            f"validiert ist.")
    para(d, f"Mit Beschluss B-06 wird ab der Sitzung 02 des Steering Committee ausschliesslich "
            f"der validierte Wert berichtet. Damit ist der in der ersten Sitzung entstandene "
            f"Widerspruch (Vorgang I-03) geschlossen.")

    d.add_heading("3.4 Pensionsverpflichtung", level=2)
    para(d, f"Das Risiko R-13 wird erneut aufgerufen. Die Pensionsverpflichtung beim "
            f"Zielunternehmen ist im Rahmen der oeffentlich bekannten Risikofaktoren der "
            f"Transaktion ausdruecklich benannt worden. Fuer die Eroeffnungsbilanz (M-04) ist "
            f"eine versicherungsmathematische Pruefung erforderlich; sie ist als OP-09 erfasst, "
            f"faellig {C.de(C.ACTIONS[8][7])}.")

    d.add_heading("4. Festlegungen", level=1)
    grid(d, ["Nr.", "Festlegung", "Begruendung", "Verantwortlich", "Wirksam ab"], [
        ["F-01", "Die Kontenzuordnung wird vor Beginn der Saldenabstimmung final freigegeben",
         "Sonst wird die Abstimmung zweimal gemacht", C.nm("WS1"), "03.10.2016"],
        ["F-02", "Der Konsolidierungskreis wird erst nach Freigabe von M-05 festgelegt",
         "Eine vorlaeufige Festlegung wuerde eine Nacharbeit im Quartalsabschluss ausloesen",
         C.nm("dach_fin"), "nach Freigabe M-05"],
        ["F-03", "Berichtet wird ausschliesslich der von Finance validierte Synergiewert",
         "Umsetzung von Beschluss B-06", C.nm("synergy"), C.de(C.STEERCO_02)],
    ], [1.2, 5.4, 5.0, 2.6, 2.8])

    d.add_heading("5. Aufgaben", level=1)
    rows = [[a[0], a[2][:70], C.nm(a[4]), C.de(a[7]), STATUS_DE[a[9]]]
            for a in C.ACTIONS if a[5] == "WS1"]
    grid(d, ["Nr.", "Aufgabe", "Verantwortlich", "Faellig am", "Status"], rows,
         [1.4, 6.8, 3.0, 2.4, 3.4])

    d.add_heading("6. Risiken zur Eskalation an das VCIO", level=1)
    rows = [[r[0], r[2][:76], f"{C.sev(r)}", r[7][:60], C.nm(r[4]),
             "ja" if r[11] else "nein"] for r in C.ws_risks("WS1")]
    grid(d, ["Nr.", "Risiko", "Schwere", "Gegenmassnahme", "Verantwortlich",
             "Eskalation"], rows, [1.2, 5.4, 1.6, 4.4, 2.6, 1.8])

    d.add_heading("7. Naechster Termin", level=1)
    kv(d, [("Termin", "04.10.2016, 14:00 Uhr MESZ"),
           ("Schwerpunkt", "Eroeffnungsbilanz M-04, Stand der versicherungsmathematischen "
                           "Pruefung")])

    disclaimer(d, german=True)
    d.save(C.OUT / "DellEMC_VCIO_Teilprojekt_Protokoll_Finanzen_2016-09-27.docx")


# =====================================================================
# 4. Expertenrunde ERP (DE)
# =====================================================================
def expertenrunde():
    d = doc()
    head(d, f"{C.PROGRAM} - Ergebnisprotokoll Expertenrunde",
         "Thema: ERP-Konsolidierung, Zielsystem und Migrationspfad  |  Ad-hoc-Fachrunde",
         [("Dokument", "Ergebnisprotokoll der Expertenrunde"),
          ("Programm", C.PROGRAM),
          ("Thema", "ERP-Konsolidierung: Zielsystem, Migrationspfad, Auswirkung auf M-07"),
          ("Anlass", "Eskalation aus dem VCIO Weekly vom 20.09.2016, Meilenstein M-07 "
                     "nicht zum Basistermin erreichbar"),
          ("Datum und Uhrzeit", f"{C.de(C.D_EXPERT)}, 13:00 bis 16:30 Uhr MESZ"),
          ("Moderation", f"{C.nm('WS3')}, {C.role('WS3')}"),
          ("Protokollfuehrung", f"{C.nm('dach_it')}, {C.role('dach_it')}"),
          ("Verteiler", "VCIO, WS3 IT, WS1 Finanzen, WS5 Lieferkette, Regionalleitung DACH"),
          ("Klassifizierung", "Streng vertraulich"),
          ("Version und Status", "v1.0, freigegeben")])

    d.add_heading("1. Zielsetzung der Runde", level=1)
    para(d, "Die Runde soll eine Empfehlung fuer das ERP-Zielsystem und den Migrationspfad "
            "herbeifuehren und die Auswirkung auf den Meilenstein M-07 sowie auf die "
            "abhaengigen Arbeitspakete in der Lieferkette bewerten. Ohne diese Entscheidung "
            "kann weder die Stammdatenbereinigung abgeschlossen noch der Konsolidierungskreis "
            "technisch abgebildet werden.")
    para(d, "Die parallel laufende Modernisierung des Altsystems beim Zielunternehmen ist in "
            "den oeffentlichen Risikofaktoren der Transaktion ausdruecklich als Risiko "
            "benannt. Genau dieses Risiko ist nun eingetreten.")

    d.add_heading("2. Teilnehmer und eingebrachte Expertise", level=1)
    grid(d, ["Name", "Rolle", "Gesellschaft", "Eingebrachte Expertise"], [
        [C.nm("WS3"), C.role("WS3"), C.NEWCO, "Anwendungsarchitektur, Programmleitung IT"],
        [C.nm("dach_it"), C.role("dach_it"), "Dell EMC DACH", "Datenmigration, Stammdaten"],
        [C.nm("dach_fin"), C.role("dach_fin"), "Dell EMC DACH", "Rechnungswesen, Kontenlogik"],
        [C.nm("WS5"), C.role("WS5"), C.NEWCO, "Lieferkette, Auftragsabwicklung"],
        [C.nm("WS1"), C.role("WS1"), C.NEWCO, "Konzernrechnungswesen"],
        [C.nm("advisor"), C.role("advisor"), C.ADVISOR,
         "Vorgehen bei ERP-Konsolidierungen in Transaktionen"],
        [C.nm("imo_mgr"), C.role("imo_mgr"), C.NEWCO, "Abhaengigkeiten, Terminplan"],
    ], [3.2, 4.6, 3.4, 5.8])

    d.add_heading("3. Ausgangslage", level=1)
    d.add_heading("3.1 Bestehende Systemlandschaft, Ausschnitt DACH", level=2)
    grid(d, ["System", "Herkunft", "Modul", "Anwender", "Wartung bis", "Kritikalitaet"], [
        ["ERP A, Produktivsystem", "Erwerber", "FI, CO, MM, SD", "3.400", "31.12.2019", "hoch"],
        ["ERP B, Produktivsystem", "Zielunternehmen", "FI, CO, MM", "1.900", "30.06.2018",
         "hoch"],
        ["ERP B, Modernisierungsprojekt", "Zielunternehmen", "FI, CO", "laufend",
         "nicht anwendbar", "hoch"],
        ["Stammdatenwerkzeug", "Zielunternehmen", "Material, Lieferant", "240", "31.03.2017",
         "mittel"],
        ["Auftragsabwicklung", "Erwerber", "SD, Logistik", "1.150", "31.12.2018", "hoch"],
        ["Berichtsschicht", "beide", "Konsolidierung", "310", "unterschiedlich", "mittel"],
    ], [3.8, 3.0, 3.0, 1.8, 2.6, 2.4])

    d.add_heading("3.2 Rahmenbedingungen und Restriktionen", level=2)
    for x in [
        "Der Meilenstein M-07 ist gate-relevant. Ohne freigegebenes Zielbild kann die "
        "Migrationsplanung nicht beginnen.",
        "Die Wartung fuer das Altsystem des Zielunternehmens laeuft am 30.06.2018 aus. Eine "
        "Verlaengerung ist moeglich, aber nicht kostenneutral.",
        "Die Annahme AS-01, dass der Wartungsvertrag mit drei Monaten Frist kuendbar ist, ist "
        "noch nicht bestaetigt. Trifft sie nicht zu, entfaellt die Lizenzsynergie S-04.",
        "Die erneute Extraktion der Materialstammdaten (A-023) ist Voraussetzung fuer jede "
        "belastbare Aufwandsschaetzung.",
        f"Die Zusage zum Wertrealisierungsbericht am {C.de(C.DAY100)} steht und wird nicht "
        f"zur Disposition gestellt.",
    ]:
        d.add_paragraph(x, style="List Bullet")

    d.add_heading("4. Diskutierte Optionen", level=1)
    grid(d, ["Option", "Beschreibung", "Aufwand (PT)", "Kosten (TEUR)", "Dauer",
             "Vorteile", "Nachteile und Risiken"], [
        ["Option 1", "Zielsystem ist ERP A des Erwerbers, Altsystem wird stillgelegt",
         "4.200", "3.150", "14 Monate",
         "Ein Zielsystem, hoechste Synergie, vorhandenes Wissen im Konzern",
         "Groesster Migrationsaufwand, hoechstes Risiko in der Lieferkette"],
        ["Option 2", "Beide Systeme bleiben, gemeinsame Berichtsschicht darueber",
         "1.400", "1.050", "6 Monate",
         "Schnell, geringes Betriebsrisiko, entkoppelt von der Modernisierung",
         "Keine Lizenzsynergie, doppelte Betriebskosten dauerhaft"],
        ["Option 3", "Zielsystem ist ERP A, Migration in zwei Wellen nach Region",
         "4.600", "3.480", "18 Monate",
         "Risiko je Welle beherrschbar, Lernkurve nutzbar",
         "Laengste Laufzeit, Synergie faellt spaeter an, zwei Umstellungsphasen"],
    ], [1.8, 4.6, 1.8, 2.0, 1.6, 3.4, 3.8])

    d.add_heading("5. Fachliche Bewertung", level=1)
    grid(d, ["Kriterium", "Gewicht", "Option 1", "Option 2", "Option 3", "Begruendung"], [
        ["Synergiebeitrag", "30 %", "5", "1", "5",
         "Nur die Stilllegung des Altsystems hebt S-04 und S-05"],
        ["Umsetzungsrisiko", "25 %", "2", "5", "4",
         "Ein Umstellungsschritt fuer alle Regionen ist das groesste Einzelrisiko"],
        ["Zeit bis zum Nutzen", "20 %", "3", "5", "2",
         "Option 2 wirkt sofort, bringt aber keinen dauerhaften Nutzen"],
        ["Betriebsstabilitaet", "15 %", "3", "4", "4",
         "Wellenmodell erlaubt Korrektur nach der ersten Welle"],
        ["Auswirkung auf M-07", "10 %", "3", "5", "3",
         "Option 2 laesst das Zielbild sofort freigeben"],
        ["Gewichtete Summe", "100 %", "3,30", "3,50", "4,00", "Option 3 fuehrt"],
    ], [3.6, 1.8, 1.8, 1.8, 1.8, 6.2])

    d.add_heading("6. Empfehlung der Expertenrunde", level=1)
    para(d, "Die Runde empfiehlt Option 3: Zielsystem ist ERP A des Erwerbers, die Migration "
            "erfolgt in zwei Wellen nach Region. Damit bleibt der volle Synergiebeitrag "
            "erhalten, waehrend das Umsetzungsrisiko je Welle beherrschbar bleibt. Der Umfang "
            "der laufenden Modernisierung des Altsystems ist einzufrieren, bis das Zielbild "
            "freigegeben ist; dies ist als Beschluss B-04 bereits gefasst.")
    para(d, f"Fuer M-07 bedeutet die Empfehlung einen Prognosetermin von "
            f"{C.de(C.MILESTONES[6][7])} statt {C.de(C.MILESTONES[6][6])}.")

    d.add_heading("6.1 Abweichende Meinung", level=2)
    para(d, f"{C.nm('WS5')} traegt ausdruecklich eine abweichende Meinung vor. Aus Sicht der "
            f"Lieferkette ist das Wellenmodell nicht risikoaermer, sondern verlaengert nur den "
            f"Zeitraum, in dem zwei Stammdatenbestaende parallel gepflegt werden muessen. "
            f"Sie haelt Option 1 fuer vorzugswuerdig, sofern die Stammdatenqualitaet vor der "
            f"Umstellung nachweislich erreicht ist.")
    para(d, f"{C.nm('advisor')} haelt fest, dass beide Positionen dieselbe Voraussetzung "
            f"teilen, naemlich eine belastbare Stammdatenqualitaet, und dass die Entscheidung "
            f"zwischen den Optionen erst nach der erneuten Extraktion (A-023) tragfaehig ist.")

    d.add_heading("7. Auswirkungen auf Meilensteine und Synergien", level=1)
    grid(d, ["Betroffen", "Teilprojekt", "Auswirkung", "Neuer Wert oder Termin", "Freigabe durch"], [
        ["M-07 Zielbild ERP", "WS3", "Verzoegerung", C.de(C.MILESTONES[6][7]),
         "Steering Committee"],
        ["A-023 Stammdatenextraktion", "WS3", "Wird kritischer Pfad",
         C.de(C.TASKS[12][5]), C.nm("WS3")],
        ["D-02 Abhaengigkeit zur Lieferkette", "WS5", "Status verzoegert",
         C.de(C.DEPENDENCIES[1][4]), "beide Teilprojektleitungen"],
        ["S-04 Lizenzsynergie", "WS3", "Zeitlich verschoben, Hoehe unveraendert",
         f"USD {C.SYNERGIES[3][7]} Mio. Ziel", C.nm("synergy")],
        ["S-05 Rechenzentrumsflaeche", "WS3", "Abhaengig von der Wellenplanung",
         f"USD {C.SYNERGIES[4][7]} Mio. Ziel", C.nm("synergy")],
    ], [4.0, 2.2, 4.0, 4.0, 3.4])

    d.add_heading("8. Aufgaben und naechste Schritte", level=1)
    grid(d, ["Nr.", "Aufgabe", "Verantwortlich", "Faellig am", "Status"], [
        ["OP-06", C.ACTIONS[5][2], C.nm("dach_it"), C.de(C.ACTIONS[5][7]),
         STATUS_DE[C.ACTIONS[5][9]]],
        ["A-023", C.TASKS[12][2], C.nm("dach_it"), C.de(C.TASKS[12][5]),
         STATUS_DE[C.TASKS[12][6]]],
        ["A-022", C.TASKS[11][2], C.nm("WS3"), C.de(C.TASKS[11][5]),
         STATUS_DE[C.TASKS[11][6]]],
        ["OP-02", C.ACTIONS[1][2], C.nm("WS3"), C.de(C.ACTIONS[1][7]),
         STATUS_DE[C.ACTIONS[1][9]]],
    ], [1.4, 6.8, 3.0, 2.4, 3.4])

    d.add_heading("9. Entscheidungsbedarf im Steering Committee", level=1)
    kv(d, [("Zu entscheiden", "Bestaetigung der Option 3 als Zielarchitektur und Freigabe des "
                              "Wiederanlaufplans fuer M-07"),
           ("Vorlage durch", f"{C.nm('WS3')}"),
           ("Erforderlich bis", C.de(C.STEERCO_02)),
           ("Konsequenz bei Verzoegerung", "Die Migrationsplanung kann nicht beginnen, die "
                                           "Abhaengigkeit D-02 bleibt verzoegert und die "
                                           "Synergien S-04 und S-05 verschieben sich in das "
                                           "Folgejahr")])

    disclaimer(d, german=True)
    d.save(C.OUT / "DellEMC_VCIO_Expertenrunde_ERP_Konsolidierung_2016-09-21.docx")


# =====================================================================
# 5. Rollenkarten (DE)
# =====================================================================
def rollenkarten():
    d = doc()
    head(d, f"{C.OFFICE} ({C.OFFICE_ABBR}) - Rollenkarten der Integrationsorganisation",
         "Mandat, Verantwortlichkeiten und Entscheidungsrechte je Rolle",
         [("Dokument", "Rollenkarten der Integrationsorganisation"),
          ("Programm", C.PROGRAM),
          ("Gueltig ab", C.de(C.D_ROLECARDS)),
          ("Erstellt durch", f"{C.nm('imo_mgr')}, {C.role('imo_mgr')}"),
          ("Freigegeben durch", f"{C.nm('prog_dir')}, {C.role('prog_dir')}"),
          ("Verteiler", "Alle Mitglieder der Integrationsorganisation"),
          ("Klassifizierung", "Streng vertraulich"),
          ("Version und Status", "v1.0, freigegeben"),
          ("Letzte Aktualisierung", f"{C.de(C.D_ROLECARDS)}, seither nicht nachgezogen")])

    d.add_heading("Hinweis zur Verwendung", level=1)
    note(d, "Je Rolle eine Karte. Die Rollenkarte beschreibt das Mandat, nicht die Person. "
            "Bei einem Personalwechsel wird ausschliesslich das Feld Rolleninhaber "
            "aktualisiert. Fuehrende Quelle fuer die aktuelle Besetzung ist die RACI-Matrix "
            "im Integration Hub. Bei Abweichungen gilt die RACI-Matrix.")

    d.add_heading("Uebergeordnete Governance", level=1)
    note(d, "Die folgende Struktur ist der oeffentlich dokumentierte Rahmen der Transaktion.")
    grid(d, ["Name", "Funktion", "Mandat in der Integration"],
         [[n, r, m] for n, r, m in C.PUBLIC_GOVERNANCE], [4.0, 6.4, 6.6])

    cards = [
        ("R01", "Programmleitung VCIO", "prog_dir",
         "Fuehrt das Integrationsbuero, verantwortet die Einhaltung der Zusagen gegenueber "
         "dem Sponsor und die Gesamtsteuerung ueber alle Teilprojekte hinweg.",
         "Freigabe von Umfangsaenderungen bis USD 5 Mio.; Eskalation an das Steering "
         "Committee darueber hinaus.",
         "Steering Committee, unverzueglich bei Gefaehrdung eines gate-relevanten "
         "Meilensteins", "1,0 FTE bis Day 100"),
        ("R02", "Integrationsmanagement VCIO", "imo_mgr",
         "Fuehrt den Terminplan, das Abhaengigkeitsregister und die Beschluss- und "
         "Massnahmenverfolgung. Stellt sicher, dass jede Uebergabe zwischen Teilprojekten "
         "einen bestaetigten Termin hat.",
         "Setzt Prognosetermine im Tracker; kann Basistermine nicht aendern.",
         "Programmleitung, woechentlich; Steering Committee ueber die Programmleitung",
         "1,0 FTE bis Day 100"),
        ("R03", "Berichtswesen VCIO", "imo_pmo",
         "Erstellt den Wochenbericht und die Unterlagen fuer das Steering Committee. "
         "Betreut den Integration Tracker als fuehrende Quelle nach Beschluss B-02.",
         "Entscheidet ueber Format und Zeitpunkt der Berichte; keine inhaltliche "
         "Entscheidungsbefugnis.",
         "Integrationsmanagement, bei Datenwiderspruechen zwischen Quellen",
         "1,0 FTE bis Day 100"),
        ("R04", "Teilprojektleitung", None,
         "Verantwortet Umfang, Termin und Qualitaet des eigenen Teilprojekts sowie die "
         "Meldung von Risiken und Abhaengigkeiten in die Register.",
         "Entscheidet innerhalb des freigegebenen Teilprojektumfangs; Abweichungen ueber "
         "10 Prozent Aufwand gehen an die Programmleitung.",
         "Programmleitung; bei Schwere ab " + str(C.ESCALATION_THRESHOLD) +
         " zusaetzlich an das Steering Committee", "0,6 bis 1,0 FTE bis Day 100"),
        ("R05", "Wertrealisierungscontrolling", "synergy",
         "Fuehrt das Synergieregister, validiert Initiativen gemeinsam mit Finance und "
         "berichtet ausschliesslich validierte Werte nach Beschluss B-06.",
         "Verweigert die Aufnahme nicht validierter Werte in die Berichterstattung.",
         "Teilprojektleitung Finanzen und Programmleitung", "0,8 FTE bis Day 100"),
        ("R06", "Risikoverantwortung Integration", "risk",
         "Fuehrt das RAID-Register, prueft die Bewertung von Wahrscheinlichkeit und "
         "Auswirkung und ueberwacht die Eskalationsschwelle.",
         "Hebt die Bewertung eines Risikos an, wenn die Faktenlage es traegt, auch gegen "
         "die Einschaetzung des Teilprojekts.",
         "Programmleitung; ab Schwere " + str(C.ESCALATION_THRESHOLD) +
         " unmittelbar an das Steering Committee", "0,6 FTE bis Day 100"),
        ("R07", "Veraenderung und Kommunikation", "change",
         "Verantwortet die Mitarbeiter- und Kundenkommunikation zur Integration, "
         "einschliesslich der digitalen Werkzeuge zu Day 1.",
         "Entscheidet ueber Kanal und Zeitpunkt; Inhalte mit rechtlicher Wirkung "
         "ausschliesslich nach Freigabe.",
         "Programmleitung; bei mitbestimmungsrelevanten Inhalten zusaetzlich an die "
         "Betriebsratskoordination", "0,8 FTE bis Day 100"),
        ("R08", "Betriebsratskoordination DACH", "br_liaison",
         "Fuehrt die Anhoerung der Betriebsraete nach den Vorschriften des "
         "Betriebsverfassungsgesetzes und haelt den Verfahrensstand fuer die "
         "Integrationsorganisation nach.",
         "Stoppt jede Veroeffentlichung, die dem Verfahrensstand vorgreift.",
         "Regionalleitung DACH und Teilprojektleitung Personal, unverzueglich",
         "0,5 FTE bis Abschluss der Anhoerung"),
    ]

    for code, name, key, mandat, rechte, esk, fte in cards:
        d.add_heading(f"Rollenkarte {code}: {name}", level=1)
        holder = f"{C.nm(key)}, {C.role(key)}" if key else \
            "je Teilprojekt besetzt, siehe Anhang"
        kv(d, [
            ("Rollenbezeichnung", name),
            ("Rolleninhaber", holder),
            ("Berichtet an", "Programmleitung VCIO" if code != "R01"
                             else f"{C.OFFICE_ABBR} Co-Leads"),
            ("Zeitanteil", fte),
            ("Mandat", mandat),
            ("Entscheidungsrechte", rechte),
            ("Eskalationsweg", esk),
            ("Pflichtberichte", "Woechentlicher Statusbeitrag in den Integration Tracker "
                                "bis Donnerstag 12:00 Uhr"),
            ("Gueltig ab", C.de(C.D_ROLECARDS)),
        ])

    d.add_heading("Anhang: Uebersicht Rollen und Besetzung", level=1)
    rows = []
    for code, key in [("R01", "prog_dir"), ("R02", "imo_mgr"), ("R03", "imo_pmo"),
                      ("R05", "synergy"), ("R06", "risk"), ("R07", "change"),
                      ("R08", "br_liaison")]:
        rows.append([code, C.role(key), C.nm(key), "Programmleitung VCIO",
                     C.de(C.D_ROLECARDS)])
    for c_ in C.WS_CODES:
        holder = C.nm("hc_prev") if c_ == "WS4" else C.nm(c_)
        rows.append([f"R04 {c_}", f"Teilprojektleitung {C.WS_NAME_DE[c_]}", holder,
                     "Programmleitung VCIO", C.de(C.D_ROLECARDS)])
    grid(d, ["Rolle", "Bezeichnung", "Rolleninhaber", "Berichtet an",
             "Letzte Aktualisierung"], rows, [2.4, 5.0, 3.4, 3.4, 2.8])
    note(d, "Achtung: Diese Uebersicht wird nicht automatisch mit der RACI-Matrix im "
            "Integration Hub synchronisiert. Sie hat den Stand der Erstfassung. Bei "
            "Abweichungen gilt die RACI-Matrix.")

    disclaimer(d, german=True)
    d.save(C.OUT / "DellEMC_VCIO_Rollenkarten_Integrationsorganisation_2016-08-01.docx")


if __name__ == "__main__":
    highlight()
    protokoll_imo()
    protokoll_ws1()
    expertenrunde()
    rollenkarten()
    print("docx done")
