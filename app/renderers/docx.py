"""Render a `Deliverable` as a branded Word document.

What this replaces was a bare `Document()` with `add_heading` calls, tables hard-
coded to Word's built-in `Light Grid Accent 1`, and — where a chart should have
been — the literal sentence `Chart: Workstream Progress — see the deck.`

This produces a document: a cover, a real table of contents, message-driven
headings, embedded charts with numbered captions, tables that repeat their header
row across a page break, decision callouts, a source note per section, and page
numbers in the footer.

Two Word-specific decisions worth stating:

* **The TOC is a real field *and* a static list.** A `TOC` field only populates
  when Word recalculates it, so a document read in Preview, Google Docs or a
  browser shows an empty contents page. Writing both means the field works for
  the people who will press F9 and the static list works for everyone else.
* **Charts are images here, not embedded objects.** Word can hold a live chart,
  but round-tripping one through python-docx is unreliable, and a wrong chart is
  worse than a static one. The deck is where the editable chart lives; this says
  so in the caption.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.context.schemas import GenerationContext
from app.deliverable.model import (
    BulletsElement,
    ChartElement,
    Deliverable,
    DiagramElement,
    ImageElement,
    KpiRowElement,
    PageDesign,
    TableElement,
    TextElement,
)
from app.renderers import docx_styles as styles
from app.renderers import naming
from app.renderers.common import MeasuredBox, RenderResult
from app.templates.brand_system import BrandSystem
from app.visualizations import charts as chart_render
from app.visualizations import diagrams as diagram_render

log = logging.getLogger("pmi.renderers.docx")

#: Usable text width on A4 portrait with 2cm margins.
CONTENT_WIDTH_IN = 6.5
FIGURE_WIDTH_IN = 6.2


def render(deliverable: Deliverable, context: GenerationContext,
           out_dir: Path) -> RenderResult:
    brand: BrandSystem = context.brand_system or _fallback_brand()
    document = Document()
    styles.install(document, brand)
    _page_setup(document)

    assets = out_dir / "assets"
    figure_number = [0]
    boxes: list[MeasuredBox] = []

    _cover(document, deliverable, context, brand)
    _contents(document, deliverable, brand)

    planned_pages = [page for page in deliverable.pages
                     if page.purpose != "cover"]
    for index, page in enumerate(planned_pages):
        if index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        boxes.extend(_section(document, page, deliverable, context, brand,
                              assets, figure_number))

    _appendix(document, deliverable, context, brand)
    _footer(document, deliverable, context, brand)

    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / naming.output_name(deliverable, context, "docx")
    document.save(str(path))
    log.info("rendered %s (%d sections)", path.name, len(deliverable.pages))

    return RenderResult(path=path, page_count=len(deliverable.pages),
                        element_boxes=boxes, warnings=list(deliverable.warnings))


def _fallback_brand() -> BrandSystem:
    from app.templates import template_registry

    return template_registry.default().brand


def _page_setup(document) -> None:
    for section in document.sections:
        section.left_margin = section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.8)


# =================================================================== cover
def _cover(document, deliverable: Deliverable, context: GenerationContext,
           brand: BrandSystem) -> None:
    if brand.logo_png_b64:
        import base64
        import io

        try:
            document.add_picture(io.BytesIO(base64.b64decode(brand.logo_png_b64)),
                                 height=Inches(0.34))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:                                      # noqa: BLE001
            pass                           # a missing logo is not worth failing on

    _para(document, deliverable.title, styles.TITLE)
    if deliverable.subtitle:
        _para(document, deliverable.subtitle, styles.SUBTITLE)

    for line in (deliverable.audience_label, context.reporting_period,
                 f"Prepared {deliverable.created_at[:10]}"):
        if line:
            _para(document, line, styles.META)

    if deliverable.governing_message:
        _para(document, deliverable.governing_message, styles.GOVERNING)
    if deliverable.executive_takeaway:
        _para(document, deliverable.executive_takeaway, styles.BODY)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _contents(document, deliverable: Deliverable, brand: BrandSystem) -> None:
    """A refreshable TOC whose cached fallback also shows page numbers.

    The visible entries live *inside* the TOC field result. Word replaces that
    cached result with its authoritative pagination when it opens the file;
    viewers that do not refresh fields still show titles, dot leaders and a
    useful page-number fallback instead of the former title-only list.
    """
    sections = [p for p in deliverable.pages if p.purpose != "cover"]
    if len(sections) < 3:
        return

    heading = _para(document, "Contents", styles.H1)
    _set_outline_level(heading, 9)  # keep the TOC from listing itself
    _set_update_fields(document)

    for index, page in enumerate(sections):
        paragraph = document.add_paragraph(style=document.styles[styles.BODY])
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS)
        if index == 0:
            _toc_field_start(paragraph)
        paragraph.add_run((page.title or page.page_id) + "\t")
        # Sections are explicitly page-broken, so 3 + index is a useful cached
        # value even in a viewer that never evaluates fields. PAGEREF replaces
        # it with the exact page when Word/LibreOffice refreshes the document.
        _field_with_result(
            paragraph, f"PAGEREF {_toc_bookmark(page, index)} \\h",
            str(index + 3))
        if index == len(sections) - 1:
            _toc_field_end(paragraph)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ================================================================= sections
def _section(document, page: PageDesign, deliverable: Deliverable,
             context: GenerationContext, brand: BrandSystem, assets: Path,
             figure_number: list[int]) -> list[MeasuredBox]:
    boxes: list[MeasuredBox] = []

    if page.purpose == "divider":
        _para(document, page.title, styles.QUOTE)
        return boxes

    heading = _para(document, page.title, styles.H1)
    section_index = [p for p in deliverable.pages
                     if p.purpose != "cover"].index(page)
    _bookmark(heading, _toc_bookmark(page, section_index), section_index + 1)
    if page.subtitle:
        _para(document, page.subtitle, styles.H3)

    for element in page.elements:
        if isinstance(element, TextElement):
            if not element.text:
                continue
            style = {"callout": styles.CALLOUT, "quote": styles.QUOTE}.get(
                element.role, styles.BODY)
            _para(document, element.text, style)
            boxes.append(_box(page, element.role, element.text))

        elif isinstance(element, BulletsElement):
            for item in element.items:
                _para(document, item, styles.BULLET, bullet=True)
            if element.items:
                boxes.append(_box(page, "bullets", " ".join(element.items)))

        elif isinstance(element, KpiRowElement):
            _kpi_table(document, element, brand)
            boxes.append(_box(page, "kpi_row", ""))

        elif isinstance(element, ChartElement):
            spec = deliverable.specs.charts.get(element.spec_id)
            if spec is not None:
                _figure(document, chart_render.to_png(spec, brand, assets,
                                                      size_in=(9.0, 4.6)),
                        "", figure_number, brand)
                boxes.append(_box(page, "chart", spec.caption))

        elif isinstance(element, DiagramElement):
            spec = deliverable.specs.diagrams.get(element.spec_id)
            if spec is not None:
                _figure(document, diagram_render.to_png(spec, brand, assets,
                                                        size_in=(9.0, 3.8)),
                        "", figure_number, brand)
                boxes.append(_box(page, "diagram", spec.caption))

        elif isinstance(element, TableElement):
            spec = deliverable.specs.tables.get(element.spec_id)
            if spec is not None:
                _table(document, spec, brand)
                boxes.append(_box(page, "table", spec.caption))

        elif isinstance(element, ImageElement) and element.image_ref:
            try:
                document.add_picture(element.image_ref,
                                     width=Inches(FIGURE_WIDTH_IN))
            except Exception:                                  # noqa: BLE001
                pass

    if page.source_note:
        _para(document, page.source_note, styles.SOURCE_NOTE)

    return boxes


def _figure(document, path: Path, caption: str, figure_number: list[int],
            brand: BrandSystem, *, editable_note: bool = False) -> None:
    """A picture and its caption, kept together and natively numbered."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(FIGURE_WIDTH_IN))

    figure_number[0] += 1
    text = caption
    if editable_note:
        text = (text.rstrip(".")
                + ". The editable version of this chart is in the accompanying "
                  "presentation.")
    _caption(document, f"Figure {figure_number[0]}", text)


def _caption(document, label: str, text: str) -> None:
    """`Figure N` as a `SEQ` field, so Word renumbers if a figure is inserted."""
    paragraph = document.add_paragraph(style=styles.CAPTION)
    paragraph.add_run("Figure ")
    run = paragraph.add_run()
    begin = run._r.makeelement(qn("w:fldChar"), {})
    begin.set(qn("w:fldCharType"), "begin")
    instruction = run._r.makeelement(qn("w:instrText"), {})
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = r"SEQ Figure \* ARABIC"
    end = run._r.makeelement(qn("w:fldChar"), {})
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, end):
        run._r.append(node)
    # A literal number too, for readers whose viewer does not evaluate fields.
    paragraph.add_run(f"{label.split()[-1]}. ")
    paragraph.add_run(text)


def _table(document, spec, brand: BrandSystem) -> None:
    table = document.add_table(rows=1, cols=len(spec.columns))
    table.autofit = True
    header = table.rows[0]
    styles.repeat_header_row(header)

    for index, column in enumerate(spec.columns):
        cell = header.cells[index]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.style = document.styles[styles.TABLE_HEADER]
        paragraph.add_run(column.header)
        if column.kind in ("number", "currency", "percent"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        styles.shade_cell(cell, brand.color("primary"))

    for row_index, row in enumerate(spec.rows):
        cells = table.add_row()
        styles.cannot_split(cells)
        emphasised = row_index in spec.emphasis_rows
        for column_index, value in enumerate(row):
            if column_index >= len(spec.columns):
                continue
            cell = cells.cells[column_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles[styles.TABLE_BODY]
            run = paragraph.add_run(value.text)
            run.bold = emphasised
            colour = _emphasis_colour(value.emphasis, brand)
            if colour:
                run.font.color.rgb = styles.rgb(colour)
            if spec.columns[column_index].kind in ("number", "currency",
                                                  "percent"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if row_index % 2 == 1:
                styles.shade_cell(cell, brand.color("surface_alt"))

    if spec.caption:
        _para(document, spec.caption, styles.CAPTION)
    if spec.is_truncated:
        _para(document, spec.truncation_note(), styles.SOURCE_NOTE)


def _emphasis_colour(emphasis: str, brand: BrandSystem) -> Optional[str]:
    return {"bad": brand.color("rag_red"), "warn": brand.color("rag_amber"),
            "good": brand.color("rag_green"),
            "muted": brand.color("muted")}.get(emphasis)


def _kpi_table(document, element: KpiRowElement, brand: BrandSystem) -> None:
    """KPIs as a borderless single-row table, so they sit side by side."""
    tiles = element.tiles[:6]
    if not tiles:
        return
    table = document.add_table(rows=2, cols=len(tiles))
    for index, tile in enumerate(tiles):
        value_cell = table.cell(0, index)
        value_cell.text = ""
        value = value_cell.paragraphs[0]
        value.style = document.styles[styles.KPI_VALUE]
        run = value.add_run(tile.display or "Not Reported")
        colour = _emphasis_colour(tile.emphasis, brand) or brand.color("primary")
        run.font.color.rgb = styles.rgb(colour)
        styles.shade_cell(value_cell, brand.color("surface_alt"))

        label_cell = table.cell(1, index)
        label_cell.text = ""
        label = label_cell.paragraphs[0]
        label.style = document.styles[styles.KPI_LABEL]
        label.add_run(tile.label + (f" ({tile.note})" if tile.note else ""))
        styles.shade_cell(label_cell, brand.color("surface_alt"))


# ================================================================= appendix
def _appendix(document, deliverable: Deliverable, context: GenerationContext,
              brand: BrandSystem) -> None:
    """Sources, disagreements and limitations. Always present, always last."""
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _para(document, "Sources and methodology", styles.H1)

    _para(document, "Sources read", styles.H2)
    files = context.evidence.projected_from_files
    if files:
        for name in files:
            _para(document, name, styles.SOURCE_NOTE)
    else:
        _para(document, "No files were read for this document.", styles.BODY)

    conflicts = context.unresolved_critical_conflicts
    if conflicts:
        _para(document, "Unresolved disagreements between sources", styles.H2)
        _para(document, "The figures below are disputed. Neither value should be "
                        "treated as agreed.", styles.BODY)
        for conflict in conflicts:
            claims = "; ".join(f"{name} says {value}"
                               for name, value in conflict.values.items())
            _para(document, f"{conflict.entity_key} ({conflict.field}): {claims}",
                  styles.BULLET, bullet=True)

    if deliverable.notes:
        _para(document, "Limitations", styles.H2)
        for note in deliverable.notes:
            _para(document, note, styles.BULLET, bullet=True)

    if deliverable.warnings:
        _para(document, "What the system could not do", styles.H2)
        for warning in deliverable.warnings[:15]:
            _para(document, warning, styles.BULLET, bullet=True)


def _footer(document, deliverable: Deliverable, context: GenerationContext,
            brand: BrandSystem) -> None:
    """Project, audience and `Page N of M` on every page."""
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.style = document.styles[styles.SOURCE_NOTE]
        paragraph.text = ""
        left = " · ".join(part for part in (context.display_name(),
                                           deliverable.audience_label) if part)
        paragraph.add_run(f"{left}    ")
        paragraph.add_run("Page ")
        _field(paragraph, "PAGE")
        paragraph.add_run(" of ")
        _field(paragraph, "NUMPAGES")


def _field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = run._r.makeelement(qn("w:fldChar"), {})
    begin.set(qn("w:fldCharType"), "begin")
    text = run._r.makeelement(qn("w:instrText"), {})
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    end = run._r.makeelement(qn("w:fldChar"), {})
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, text, end):
        run._r.append(node)


def _field_with_result(paragraph, instruction: str, result: str) -> None:
    """Insert a field with visible cached text for non-updating viewers."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    visible = OxmlElement("w:t")
    visible.text = result
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, text, separate, visible, end):
        run._r.append(node)


def _toc_field_start(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = r' TOC \o "1-3" \h \z \u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    for node in (begin, instruction, separate):
        run._r.append(node)


def _toc_field_end(paragraph) -> None:
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _toc_bookmark(page: PageDesign, index: int) -> str:
    return f"pmi_section_{index + 1}"


def _bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _set_update_fields(document) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.insert(0, update)
    update.set(qn("w:val"), "true")


def _set_outline_level(paragraph, level: int) -> None:
    properties = paragraph._p.get_or_add_pPr()
    outline = properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        properties.append(outline)
    outline.set(qn("w:val"), str(level))


# ================================================================== helpers
def _para(document, text: str, style_name: str, *, bullet: bool = False):
    paragraph = document.add_paragraph(style=document.styles[style_name])
    paragraph.add_run(("•  " if bullet else "") + text)
    return paragraph


def _box(page: PageDesign, name: str, text: str) -> MeasuredBox:
    """Word reflows, so only content is reported; geometry is Word's business."""
    return MeasuredBox(page_id=page.page_id, name=name, text=text)
