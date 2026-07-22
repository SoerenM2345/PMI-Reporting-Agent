"""Word rendering. `python-docx` is already a dependency (the Word *extractor*
uses it), so this format costs nothing to add.

Worth having because a Word file is the one deliverable a PMI consultant can
edit in front of a client. A deck is for presenting and a workbook is for
checking; a document is for redlining, and "send me that in Word so I can mark
it up" is the most common thing anyone says about a status report.
"""
from __future__ import annotations

import logging
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.report.content import ReportContent, Section

log = logging.getLogger("pmi.render.docx")

RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xF9, 0xA8, 0x25)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREY = RGBColor(0x75, 0x75, 0x75)

_COLOUR = {"bad": RED, "warn": AMBER, "good": GREEN, "muted": GREY}


def render(content: ReportContent, out_dir: Path) -> Path:
    document = Document()

    document.add_heading(content.title, level=0)
    if content.subtitle:
        _muted(document.add_paragraph(content.subtitle))
    if content.meta_line:
        _muted(document.add_paragraph(" · ".join(content.meta_line)))

    for section in content.narrative():
        _section(document, section, content)

    if content.footer:
        document.add_paragraph()
        _muted(document.add_paragraph(content.footer))

    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"PMI_Report_{content.audience.value}_{date.today().isoformat()}.docx"
    document.save(str(path))
    log.info("wrote %s", path.name)
    return path


def _section(document, section: Section, content: ReportContent) -> None:
    # The management message is the heading; the section name sits under it in
    # smaller type, exactly as on the slide (§12.5).
    document.add_heading(section.headline, level=1)
    _muted(document.add_paragraph(section.label))

    rendered = False
    for block in section.blocks:
        if _block(document, block, content):
            rendered = True

    if not rendered and section.empty_explanation:
        _muted(document.add_paragraph(section.empty_explanation))


def _block(document, block, content: ReportContent) -> bool:
    if block.kind == "prose":
        if block.title:
            document.add_paragraph(block.title, style="Intense Quote")
        document.add_paragraph(block.text)
        return True

    if block.kind == "bullets":
        if not block.items:
            return False
        if block.title:
            run = document.add_paragraph().add_run(block.title)
            run.bold = True
        for item in block.items:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(item.text)
            if item.emphasis in _COLOUR:
                run.font.color.rgb = _COLOUR[item.emphasis]
        return True

    if block.kind == "tiles":
        if not block.tiles:
            return False
        table = document.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for tile in block.tiles:
            value = (content.facts.display(tile.fact_key)
                     if tile.fact_key else (tile.value or "Not Reported"))
            cells = table.add_row().cells
            cells[0].text = tile.label
            run = cells[1].paragraphs[0].add_run(value)
            run.bold = True
            if tile.emphasis in _COLOUR:
                run.font.color.rgb = _COLOUR[tile.emphasis]
        return True

    if block.kind == "table":
        return _table(document, block)

    if block.kind == "chart":
        caption = block.caption or block.builder.replace("_", " ").title()
        _muted(document.add_paragraph(f"Chart: {caption} — see the deck."))
        return True

    return False


def _table(document, block) -> bool:
    if block.is_derived:
        entity = block.query.entity if block.query else "rows"
        _muted(document.add_paragraph(
            f"Full {entity} listing — included in the Excel dashboard."
        ))
        return True
    if not block.rows:
        return False

    if block.title:
        run = document.add_paragraph().add_run(block.title)
        run.bold = True

    table = document.add_table(rows=1, cols=len(block.columns))
    table.style = "Light Grid Accent 1"
    for index, column in enumerate(block.columns):
        run = table.rows[0].cells[index].paragraphs[0].add_run(column.header)
        run.bold = True

    limit = block.row_limit or len(block.rows)
    for row in block.rows[:limit]:
        cells = table.add_row().cells
        for index, cell in enumerate(row[:len(block.columns)]):
            run = cells[index].paragraphs[0].add_run(cell.text)
            run.font.size = Pt(9)
            if cell.emphasis in _COLOUR:
                run.font.color.rgb = _COLOUR[cell.emphasis]

    if block.note:
        _muted(document.add_paragraph(block.note))
    return True


def _muted(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GREY
    return paragraph
