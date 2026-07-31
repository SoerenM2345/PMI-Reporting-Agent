"""Word extractor (spec §5.4): meeting minutes, workstream reports, decision logs.

Tracks the current heading while walking the document, so a record extracted from a
paragraph under "3. Risks and Issues" carries that section in its provenance. That
matters for a 20-page steering pack, where "somewhere in minutes.docx" is not a
citation anybody can act on.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.extractors.base import (
    classify_table,
    extract_actions_from_text,
    find_progress_mentions,
    make_source,
    rows_to_records,
)
from app.models.pmi import ExtractionMethod, SourceFormat

suffixes: tuple[str, ...] = (".docx",)
format: SourceFormat = SourceFormat.WORD

_DECISION_PREFIXES = ("decision:", "decided:", "beschluss:", "entscheidung:")


def extract(path: Path) -> list[dict]:
    document = Document(str(path))
    records: list[dict] = []

    section = ""
    table_index = 0
    text_parts: list[str] = []

    for block in _iter_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue

            if block.style.name.startswith("Heading"):
                section = text
                continue

            text_parts.append(text)

            # Minutes record decisions as prose; §6.8 wants them as entities.
            lowered = text.casefold()
            if any(lowered.startswith(p) for p in _DECISION_PREFIXES):
                records.append({
                    "type": "decision",
                    "title": text.split(":", 1)[1].strip() or text,
                    "decision_body": section or "Steering Committee",
                    "source": make_source(
                        path.name, format,
                        section_name=section or None,
                        original_value=text,
                        extraction_method=ExtractionMethod.TEXT_REGEX,
                    ),
                })

        elif isinstance(block, Table):
            table_index += 1
            grid = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            if len(grid) < 2:
                continue
            source = make_source(
                path.name, format,
                section_name=section or None,
                table_name=f"table {table_index}",
                extraction_method=ExtractionMethod.TABLE_PARSE,
            )
            records.extend(rows_to_records(
                grid[0], grid[1:],
                classify_table(grid[0], context=section),
                source,
            ))

    text = "\n".join(text_parts)
    source = make_source(path.name, format,
                         extraction_method=ExtractionMethod.TEXT_REGEX)

    for pct in find_progress_mentions(text):
        records.append({"type": "kpi", "name": "Overall Progress", "value": pct,
                        "unit": "%", "source": source})
    for item in extract_actions_from_text(text):
        records.append({**item, "source": source})
    if text.strip():
        records.append({"type": "note", "text": text.strip()[:4000], "source": source})

    return records


def _iter_blocks(document: _Document):
    """Yield paragraphs and tables in document order.

    python-docx exposes `.paragraphs` and `.tables` as separate lists, which loses
    the interleaving — and with it the ability to know which heading a table sits
    under. Walking the XML body preserves it.
    """
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)
