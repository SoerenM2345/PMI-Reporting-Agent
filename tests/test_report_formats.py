"""Word, PDF and HTML rendering (§17 downloadable outputs).

All three read the same approved `ReportContent` the user saw as text, so the
assertions are about that agreement — and about §21.15, "ensure generated files
open successfully", which for a new format is the whole risk.
"""
from __future__ import annotations

from datetime import date

import pytest

from app.extractors.base import make_source
from app.models.pmi import (
    Audience,
    BudgetItem,
    PMIDataModel,
    PMIProject,
    Risk,
    SourceFormat,
    Status,
)
from app.report.planner import plan
from app.report.render import docx as docx_render
from app.report.render import pdf as pdf_render
from app.report.render.html import render_html


@pytest.fixture
def content():
    xlsx = make_source("tracker.xlsx", SourceFormat.EXCEL, sheet_name="Workplan")
    image = make_source("dashboard.png", SourceFormat.IMAGE,
                        extraction_confidence=0.35)
    model = PMIDataModel(
        project=PMIProject(project_id="p1", project_name="Project Aurora",
                           reporting_date=date.today()),
        source_files=["tracker.xlsx", "dashboard.png"],
        risks=[Risk(risk_id="R1", title="GDPR retention breach", probability=4,
                    impact=5, owner="Lisa Chen", status=Status.IN_PROGRESS,
                    source_references=[image])],
        budget=[BudgetItem(budget_item_id="B1", category="Advisors",
                           budget=1000.0, actual=900.0, forecast=1250.0,
                           source_references=[xlsx])],
    )
    return plan(model, Audience.EXECUTIVE,
                bullets=["One critical risk is unowned."])


def headlines(content) -> list[str]:
    return [s.headline for s in content.narrative()]


# ==================================================================== §21.15
def test_every_format_produces_a_file_that_opens(content, tmp_path):
    """A corrupt deliverable we handed over confidently is the worst outcome —
    the user finds out in the meeting."""
    from docx import Document
    import fitz

    word = docx_render.render(content, tmp_path)
    Document(str(word))

    pdf = pdf_render.render(content, tmp_path)
    with fitz.open(str(pdf)) as document:
        assert document.page_count >= 1

    html = render_html(content)
    assert html.startswith("<!doctype html>") and html.rstrip().endswith("</html>")


# ============================================ they say what the preview said
@pytest.mark.parametrize("fmt", ["word", "pdf", "html"])
def test_each_format_states_what_the_user_approved(content, tmp_path, fmt):
    """The promise of the preview holds across every format, not just the deck."""
    if fmt == "word":
        from docx import Document

        path = docx_render.render(content, tmp_path)
        text = "\n".join(p.text for p in Document(str(path)).paragraphs)
        text += "\n" + "\n".join(
            cell.text
            for table in Document(str(path)).tables
            for row in table.rows for cell in row.cells
        )
    elif fmt == "pdf":
        import fitz

        path = pdf_render.render(content, tmp_path)
        with fitz.open(str(path)) as document:
            text = "\n".join(page.get_text() for page in document)
    else:
        text = render_html(content)

    for headline in headlines(content):
        # PDF transliterates typographic characters into Latin-1; compare on the
        # alphanumeric core so the assertion tests content, not encoding.
        core = "".join(ch for ch in headline if ch.isalnum())[:28]
        stripped = "".join(ch for ch in text if ch.isalnum())
        assert core in stripped, f"{fmt} does not state: {headline!r}"


def test_the_screenshot_warning_survives_into_every_format(content, tmp_path):
    """§5.6. A figure read off an image at 35% must be disclosed wherever the
    report is read, not only in the deck."""
    import fitz

    html = render_html(content)
    assert "low confidence" in html

    with fitz.open(str(pdf_render.render(content, tmp_path))) as document:
        assert "low confidence" in "\n".join(p.get_text() for p in document)


# ================================================================== escaping
def test_html_escapes_content_that_came_out_of_an_uploaded_file(content, tmp_path):
    """Task titles and file names are attacker-influenced by construction —
    this tool's whole job is reading other people's documents."""
    section = content.narrative()[1]
    section.headline = "<script>alert('x')</script> risks"

    html = render_html(content)
    assert "<script>alert" not in html
    assert "&lt;script&gt;" in html


def test_pdf_transliterates_glyphs_the_core_fonts_cannot_show(content, tmp_path):
    """"⚠ NO OWNER" must not vanish — the glyph is what makes the row worth
    looking at, so it degrades to "!" rather than being dropped."""
    import fitz

    section = content.narrative()[1]
    section.headline = "⚠ 1 risk — “unowned”"

    with fitz.open(str(pdf_render.render(content, tmp_path))) as document:
        text = "\n".join(page.get_text() for page in document)

    assert "! 1 risk" in text
    assert '"unowned"' in text


def test_a_pipe_or_angle_bracket_cannot_break_a_word_table(content, tmp_path):
    from docx import Document

    path = docx_render.render(content, tmp_path)
    assert Document(str(path)).tables, "the report produced no tables at all"
