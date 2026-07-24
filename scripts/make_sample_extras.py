"""The rest of the §19 sample set: synergy tracker, PDF report, HTML export,
workstream status report, and a CSV milestone tracker.

§19 asks that the sample project contain *intentional* inconsistencies, and it does.
Every one of them is a disagreement a real integration produces, and each trips a
different check:

  * Synergy realized: 12.4 MEUR in the tracker, 6.5 MEUR in the PDF   → PMI-008 (critical)
  * ERP go-live: 15-09 in the masterplan, 30-09 on the whiteboard      → PMI-006 (critical)
  * Overall progress: 82% in Excel, 75% in the deck                    → PMI-002 (critical)
  * Day 1 readiness: 92% in the CSV, 85% in the HTML export            → PMI-010 (critical)
  * A risk with no mitigation owner                                    → COMP-003
  * A budget line with no forecast                                     → COMP-006

Run:  python scripts/make_sample_extras.py
"""
from __future__ import annotations

from pathlib import Path

SAMPLES = Path(__file__).resolve().parents[1] / "data" / "samples"
SAMPLES.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------- synergy tracker
def make_synergy_tracker() -> Path:
    import xlsxwriter

    path = SAMPLES / "synergy_tracker.xlsx"
    workbook = xlsxwriter.Workbook(str(path))
    header = workbook.add_format({"bold": True, "bg_color": "#2E7D32",
                                  "font_color": "white"})

    sheet = workbook.add_worksheet("Synergies")
    sheet.write_row(0, 0, [
        "Synergy", "Type", "Workstream", "Owner", "Target", "Realized", "Forecast",
        "Planned Realization Date", "Confidence", "Status",
    ], header)

    rows = [
        # 12.4 realized here; the PDF report says 6.5. That is PMI-008, and §9 makes
        # any synergy disagreement critical — it is the number the deal was justified with.
        ["Procurement consolidation", "Cost synergy", "Procurement", "Marco Rossi",
         18_000_000, 12_400_000, 16_000_000, "2026-12-31", "High", "In Progress"],
        ["IT license rationalization", "Cost synergy", "IT", "Jonas Weber",
         6_000_000, 2_100_000, 5_200_000, "2027-03-31", "Medium", "In Progress"],
        ["Cross-sell to target customer base", "Revenue synergy", "Sales", "Sofia Ivanova",
         9_000_000, 800_000, 4_500_000, "2027-06-30", "Low", "In Progress"],
        # No realization date -> COMP-007.
        ["Real estate footprint reduction", "Cost synergy", "Real Estate", "Anna Schmidt",
         4_000_000, 0, 3_000_000, "", "Medium", "Not Started"],
    ]
    for index, row in enumerate(rows, start=1):
        sheet.write_row(index, 0, row)

    budget = workbook.add_worksheet("Integration Budget")
    budget.write_row(0, 0, ["Category", "Budget", "Actual", "Committed", "Forecast",
                            "Currency"], header)
    budget_rows = [
        ["External advisors", 500_000, 420_000, 60_000, 520_000, "EUR"],
        ["System migration", 1_200_000, 950_000, 180_000, 1_310_000, "EUR"],
        ["Retention", 800_000, 610_000, 90_000, 780_000, "EUR"],
        # Total does NOT equal the sum of the lines above (2,610,000) -> MATH-010.
        ["Total", 2_500_000, 1_980_000, 330_000, 2_610_000, "EUR"],
    ]
    for index, row in enumerate(budget_rows, start=1):
        budget.write_row(index, 0, row)

    workbook.close()
    return path


# ------------------------------------------------------------------ CSV milestones
def make_csv() -> Path:
    path = SAMPLES / "milestone_tracker.csv"
    path.write_text(
        "Project Aurora — Milestone Tracker (exported 01-07-2026)\n"
        "\n"
        "Milestone,Owner,Workstream,Due Date,Status,Progress %\n"
        "Legal close,Anna Schmidt,Legal,01-06-2026,Done,100\n"
        "Day 1 readiness,Anna Schmidt,Day 1 Readiness,15-06-2026,Done,100\n"
        "Payroll migration,Lisa Chen,HR,31-07-2026,In Progress,70\n"
        # 15-09 here; the whiteboard photo says 30-09 -> PMI-006, critical (go-live).
        "ERP go-live,Jonas Weber,IT,15-09-2026,In Progress,45\n"
        "TSA exit,Marco Rossi,TSA,31-12-2026,Not Started,0\n"
        "\n"
        "Day 1 readiness: 92%\n",
        encoding="utf-8",
    )
    return path


# ---------------------------------------------------------------- HTML dashboard
def make_html() -> Path:
    path = SAMPLES / "portal_dashboard_export.html"
    path.write_text(
        """<!doctype html>
<html><head><meta charset="utf-8"><title>Project Aurora — Portal Export</title>
<style>body{font-family:Arial}table{border-collapse:collapse}td,th{border:1px solid #ccc;padding:6px}</style>
</head><body>
<h1>Project Aurora — Integration Portal</h1>
<p>Exported from the PMO portal on 01-07-2026.</p>

<!-- 85% here; the CSV tracker says 92%. That is PMI-010, and §9 makes any Day 1
     readiness disagreement critical: on Day 1 the company either can pay its
     people or it cannot. -->
<p>Overall progress: 78%<br>Day 1 readiness: 85%</p>

<h2>Open issues</h2>
<table>
<caption>Issue log</caption>
<tr><th>Issue</th><th>Owner</th><th>Severity</th><th>Status</th><th>Due Date</th></tr>
<tr><td>Duplicate vendor master records blocking PO creation</td><td>Marco Rossi</td>
    <td>High</td><td>Open</td><td>2026-07-20</td></tr>
<tr><td>VPN capacity insufficient for merged headcount</td><td>Jonas Weber</td>
    <td>Medium</td><td>In Progress</td><td>2026-07-31</td></tr>
<tr><td>Works council consultation not yet scheduled</td><td></td>
    <td>Critical</td><td>Open</td><td>2026-07-10</td></tr>
</table>

<h2>Cross-workstream dependencies</h2>
<table>
<caption>Dependency log</caption>
<tr><th>Description</th><th>Providing Workstream</th><th>Receiving Workstream</th>
    <th>Owner</th><th>Required Date</th><th>Status</th></tr>
<tr><td>Final org structure needed to configure payroll</td><td>Organization</td>
    <td>HR</td><td>Anna Schmidt</td><td>2026-07-15</td><td>In Progress</td></tr>
<tr><td>Chart of accounts sign-off needed before ERP cutover</td><td>Finance</td>
    <td>IT</td><td></td><td>2026-08-01</td><td>Open</td></tr>
</table>

<p>Action: confirm works council timing with Legal -> Anna Schmidt</p>
</body></html>
""",
        encoding="utf-8",
    )
    return path


# ------------------------------------------------------------------- PDF report
def make_pdf() -> Path:
    """A text PDF (not a scan) — the scanned path is exercised by the image pipeline."""
    import fitz  # PyMuPDF

    path = SAMPLES / "steerco_pack.pdf"
    doc = fitz.open()

    page = doc.new_page()
    y = 60
    lines = [
        ("Project Aurora — Steering Committee Pack", 18, True),
        ("Reporting date: 01-07-2026", 10, False),
        ("", 10, False),
        ("Management summary", 14, True),
        ("The integration remains broadly on track. Overall progress is reported at 75%.", 10, False),
        ("Two critical risks require Steering Committee attention this period.", 10, False),
        ("", 10, False),
        ("Synergies", 14, True),
        # 6.5 MEUR here; the synergy tracker says 12.4 MEUR -> PMI-008 (critical).
        ("Synergies captured to date: EUR 6.5 million against a target of EUR 37 million.", 10, False),
        ("Procurement consolidation remains the largest single contributor.", 10, False),
        ("", 10, False),
        ("Decisions required", 14, True),
        ("Decision: Approve the additional EUR 300k for ERP cutover contingency.", 10, False),
        ("Decision: Approve the revised TSA exit date of 31-12-2026.", 10, False),
        ("", 10, False),
        ("Actions", 14, True),
        ("Action: Confirm works council consultation date -> Anna Schmidt", 10, False),
        ("Action: Re-baseline the ERP cutover plan -> Jonas Weber", 10, False),
    ]
    for text, size, bold in lines:
        if text:
            page.insert_text(
                (60, y), text, fontsize=size,
                fontname="helv" if not bold else "hebo",
            )
        y += size + 8

    doc.save(str(path))
    doc.close()
    return path


# ------------------------------------------------- workstream status report (Word)
def make_workstream_report() -> Path:
    from docx import Document

    path = SAMPLES / "workstream_status_it.docx"
    doc = Document()

    doc.add_heading("IT Workstream — Weekly Status Report", level=1)
    doc.add_paragraph("Workstream: Information Technology")
    doc.add_paragraph("Lead: Jonas Weber")
    doc.add_paragraph("Reporting period: week 27, 2026")

    doc.add_heading("Progress", level=2)
    doc.add_paragraph(
        "Overall progress: 41%. The ERP cutover remains the critical path item."
    )

    doc.add_heading("Achievements", level=2)
    doc.add_paragraph("Completed the network integration design.")
    doc.add_paragraph("Migrated 14 of 40 systems to the target estate.")

    doc.add_heading("Risks and Issues", level=2)
    doc.add_paragraph(
        # Same risk as the Excel register, but rated Medium there and High here.
        # -> PMI-007 (risk rating conflict).
        "Risk: ERP cutover slips past Q3. Impact: High. Owner: Jonas Weber."
    )

    doc.add_heading("Decisions", level=2)
    doc.add_paragraph(
        "Decision: Approve parallel-run fallback for the ERP cutover weekend."
    )

    doc.add_heading("Actions", level=2)
    doc.add_paragraph("Action: Re-baseline the cutover plan -> Jonas Weber")
    doc.add_paragraph("Action: Confirm VPN capacity uplift with the vendor -> Marco Rossi")

    doc.add_heading("Support needed", level=2)
    doc.add_paragraph(
        "Finance must sign off the chart of accounts before the ERP cutover can proceed."
    )

    doc.save(str(path))
    return path


if __name__ == "__main__":
    for make in (make_synergy_tracker, make_csv, make_html, make_pdf,
                 make_workstream_report):
        print(f"wrote {make()}")
