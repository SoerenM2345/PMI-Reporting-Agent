"""Verbatim transcript of Steering Committee session 01, 15 September 2016.

The transcript is the SOURCE; the signed minutes of 22 September 2016 are the derived
record. Every item in the minutes traces to a passage here, and nothing here contradicts
the minutes. Verified programmatically by audit.py section 8.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_TAB_ALIGNMENT

import case as C

GREEN = RGBColor.from_string(C.D_DARK)
GREY = RGBColor.from_string(C.D_GREY)
BLACK = RGBColor.from_string(C.D_BLACK)


def doc():
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Arial"
    st.font.size = Pt(9.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.12
    for s in d.sections:
        s.top_margin = Cm(1.9)
        s.bottom_margin = Cm(1.9)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    for lvl, sz in ((1, 13), (2, 10.5)):
        h = d.styles[f"Heading {lvl}"]
        h.font.name = "Arial"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.color.rgb = GREEN
    return d


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexcolor)
    tcPr.append(el)


def kv(d, pairs):
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for k, v in pairs:
        row = t.add_row()
        row.cells[0].width = Cm(4.6)
        row.cells[1].width = Cm(12.4)
        shade(row.cells[0], C.D_PALE)
        r = row.cells[0].paragraphs[0].add_run(str(k))
        r.bold = True
        r.font.size = Pt(8.5)
        r2 = row.cells[1].paragraphs[0].add_run(str(v))
        r2.font.size = Pt(8.5)
    d.add_paragraph()
    return t


def speaker_table(d, rows):
    t = d.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for j, h in enumerate(["Speaker label", "Name", "Role", "Organisation"]):
        c = t.rows[0].cells[j]
        shade(c, C.D_DARK)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
    for row in rows:
        tr = t.add_row()
        for j, v in enumerate(row):
            r = tr.cells[j].paragraphs[0].add_run(str(v))
            r.font.size = Pt(8)
    d.add_paragraph()
    return t


def turn(d, ts, speaker, text):
    """One transcript turn: [timestamp] SPEAKER: text, hanging indent."""
    p = d.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(2.7)
    pf.first_line_indent = Cm(-2.7)
    pf.space_after = Pt(5)
    r = p.add_run(f"[{ts}]  ")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    r = p.add_run(f"{speaker}: ")
    r.bold = True
    r.font.size = Pt(9.5)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    return p


def stage(d, ts, text):
    """Non-speech event line."""
    p = d.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(2.7)
    pf.first_line_indent = Cm(-2.7)
    pf.space_after = Pt(5)
    r = p.add_run(f"[{ts}]  ")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    r = p.add_run(f"[{text}]")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def note(d, text):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = GREY


# ---------------------------------------------------------------- speakers
S = {
    "AV": ("prog_dir", "Chair"),
    "DO": ("imo_mgr", ""),
    "SL": ("imo_pmo", ""),
    "MH": ("synergy", ""),
    "PN": ("risk", ""),
    "RW": ("WS1", ""),
    "ND": ("WS2", ""),
    "TB": ("WS3", ""),
    "VO": ("hc_prev", ""),
    "KM": ("WS5", ""),
    "LK": ("WS6", ""),
    "HW": ("dach_lead", ""),
    "CF": ("advisor", ""),
    "EB": ("change", ""),
}


def sp(code):
    return C.nm(S[code][0])


def build():
    d = doc()

    # ---------------------------------------------------------------- header
    p = d.add_paragraph()
    r = p.add_run(f"{C.OFFICE} ({C.OFFICE_ABBR}) - Steering Committee Session 01")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = GREEN
    p = d.add_paragraph()
    r = p.add_run("Verbatim transcript of the meeting recording")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    d.add_paragraph()

    kv(d, [
        ("Document", "Verbatim transcript, Steering Committee session 01"),
        ("Programme", f"{C.PROGRAM}, {C.OFFICE} ({C.OFFICE_ABBR})"),
        ("Parties", f"{C.ACQUIRER} and {C.TARGET}, combined as {C.NEWCO}"),
        ("Meeting", f"Steering Committee, session 01, {C.en(C.STEERCO_01)}, "
                    f"day {(C.STEERCO_01 - C.DAY1).days} after Day 1"),
        ("Recording start / end", "14:00:00 / 15:32:41 CEST"),
        ("Audio duration", "01:32:41"),
        ("Transcription", "Automated speech-to-text with human review by "
                          f"{C.nm('imo_pmo')}, {C.role('imo_pmo')}"),
        ("Transcription completed", C.en(C.STEERCO_01 + (C.STEERCO_01 - C.STEERCO_01)) +
                                    ", 18:40 CEST"),
        ("Speakers identified", f"{len(S)} of {len(S)}"),
        ("Mean recognition confidence", "0.94; passages below 0.70 marked [inaudible]"),
        ("Related record", f"Steering Committee minutes, session 01, approved and signed "
                           f"{C.en(C.STEERCO_01_SIGNED)}. The minutes are the binding record; "
                           f"this transcript is the source they were drawn from"),
        ("Retention", "Recording deleted after transcription per programme record policy; "
                      "this transcript is the retained artefact"),
        ("Classification", "Strictly confidential"),
        ("Version and status", "v1.0, reviewed and released"),
    ])

    d.add_heading("Speakers", level=1)
    rows = []
    for code, (key, extra) in S.items():
        present = "Chair" if extra else ""
        rows.append([code, C.nm(key), C.role(key) + (f" ({present})" if present else ""),
                     C.ADVISOR if key == "advisor" else C.NEWCO])
    speaker_table(d, rows)
    note(d, f"{C.nm('WS7')}, {C.role('WS7')}, sent apologies and was represented by "
            f"{C.nm('imo_mgr')}. He does not appear as a speaker. Speaker labels are the "
            f"initials assigned by the transcription tool and were verified against the "
            f"attendance sheet.")

    d.add_heading("Conventions", level=1)
    note(d, "Speech is reproduced as spoken, including hesitation, self-correction and "
            "incomplete sentences. Square brackets mark non-speech events, overlapping "
            "speech and passages the tool could not resolve. Timestamps are relative to the "
            "start of the recording. Nothing has been added, reordered or tidied.")

    # ================================================================ 1
    d.add_heading("00:00 - Opening, quorum and agenda", level=1)
    stage(d, "00:00:04", "recording begins mid-sentence, participants still joining")
    turn(d, "00:00:11", sp("AV"),
         "... so we will start on time, because we have a lot to get through and I would "
         "rather finish early than run over. Good afternoon everybody. This is the first "
         f"Steering Committee of the {C.OFFICE}, session zero one, {C.en(C.STEERCO_01)}. "
         f"We are eight days after Day 1.")
    turn(d, "00:00:34", sp("AV"),
         f"{C.nm('WS7')} has sent apologies, he is on site in Frankfurt today. "
         f"{C.nm('imo_mgr')} will represent Real Estate. Everyone else is here. That gives "
         "us all seven voting members present or represented, so we are quorate.")
    turn(d, "00:00:52", sp("DO"), "Confirmed.")
    turn(d, "00:00:54", sp("AV"),
         "Agenda went out on Monday. Any changes, additions, anything anyone wants to bring "
         "forward? ... No. Then the agenda is approved as circulated.")
    stage(d, "00:01:09", "pause, papers")
    turn(d, "00:01:14", sp("AV"),
         "One thing before we start on content. This committee is not a status meeting. "
         "I get status in writing. What I want from this room is decisions. If a workstream "
         "needs something from me or from this group, say it out loud, because if it is not "
         "said here it does not exist.")

    # ================================================================ 2
    d.add_heading("00:02 - Day 1 outcome", level=1)
    turn(d, "00:02:03", sp("AV"),
         f"{C.nm('imo_pmo')}, take us through where we actually are.")
    turn(d, "00:02:08", sp("SL"),
         "Thank you. Starting with Day 1 itself, because it is worth thirty seconds. All "
         "three Day 1 milestones landed on the day. M zero one, legal close and the brand "
         "launch, executed on the seventh. M zero two, the employee welcome guide, live for "
         f"all {C.EMPLOYEES:,} people worldwide. M zero three, the sales runbook, live for "
         f"more than {C.SALES_PROFESSIONALS:,} sellers with the cross-sell guidance for the "
         "combined portfolio.")
    turn(d, "00:02:41", sp("AV"),
         "Which is not nothing. A lot of transactions this size do not have all three of "
         "those on the day. I want that recorded, and I want it said to the teams.")
    turn(d, "00:02:53", sp("EB"),
         "It has gone out in the Friday note already. Adoption tracking on the guide runs "
         "to Day 30, so we will have a number, not just a launch.")
    turn(d, "00:03:04", sp("AV"), "Good. Carry on.")

    # ================================================================ 3
    d.add_heading("00:03 - Status by workstream", level=1)
    turn(d, "00:03:09", sp("SL"),
         "By workstream, and these are the ratings the leads submitted, not mine. "
         + ", ".join(f"{C.ws_full(c)} {C.RAG_LAST_WEEK[c].lower()}"
                     for c in C.WS_CODES) + ".")
    turn(d, "00:03:38", sp("AV"),
         "Four green, three amber. Let us take the ambers. Tax first.")
    turn(d, "00:03:45", sp("ND"),
         "Amber because the entity chain review in EMEA is running slower than I planned "
         "for. It is not a resourcing problem, it is that every jurisdiction in the chain "
         "wants its own confirmation and they do not run in parallel as cleanly as the plan "
         "assumed. Until that review closes I cannot put the rationalisation plan up for "
         "approval, and until that plan is approved Finance cannot fix the consolidation "
         "scope.")
    turn(d, "00:04:18", sp("RW"),
         "Which is D zero one in the dependency register, and to be direct about it, that "
         "one is still only confirmed from one side.")
    turn(d, "00:04:27", sp("AV"), "Meaning?")
    turn(d, "00:04:29", sp("RW"),
         "Meaning Tax has a date in the register and Finance has not agreed it, because I "
         "do not believe it yet.")
    turn(d, "00:04:36", sp("ND"), "That is fair. I would not believe it either.")
    stage(d, "00:04:39", "laughter")
    turn(d, "00:04:44", sp("AV"), "We will come back to that. IT.")
    turn(d, "00:04:48", sp("TB"),
         "Amber, and I want to be precise about why, because it is not the integration work. "
         "The enterprise system modernisation programme that was already running on the "
         "target side did not stop at close. It could not, it was ordinary course of "
         "business under the covenants. So we now have two programmes that need the same "
         "architecture decisions and, more painfully, the same eight or nine people.")
    turn(d, "00:05:19", sp("AV"),
         "That was in the risk factors of the filing. It was named. So nobody in this room "
         "gets to be surprised by it.")
    turn(d, "00:05:27", sp("TB"),
         "No. It was named as a risk and it has now happened. It is R zero one in our "
         "register, and it is the highest severity item we have.")
    turn(d, "00:05:37", sp("AV"),
         "Do you have a date for the consolidation blueprint?")
    turn(d, "00:05:41", sp("TB"),
         "Not one I would stand behind in this room. I have a date I could say to make you "
         "feel better and I am not going to say it.")
    turn(d, "00:05:50", sp("AV"),
         "Then do not. But I need it before session two, with the recovery plan behind it. "
         "Dated forecast, not a direction of travel.")
    turn(d, "00:05:59", sp("TB"), "Understood. Before session two.")
    turn(d, "00:06:03", sp("AV"),
         "Human Capital.")
    turn(d, "00:06:06", sp("VO"),
         "Amber, and it is the consultation in the German-speaking region. This is a legal "
         "gate, not a communications preference, and I want the committee to understand the "
         "difference because it changes what you can and cannot ask me to do. We cannot "
         "publish the target structure before the consultation concludes.")
    turn(d, "00:06:31", sp("HW"),
         "And the works council is not being obstructive. They are following the process "
         "the law gives them. Pushing on it does not make it faster, it makes it slower.")
    turn(d, "00:06:42", sp("AV"),
         "Do we have the timetable in writing?")
    turn(d, "00:06:45", sp("HW"),
         f"Not yet in writing. {C.nm('br_liaison')} is chasing it. I would rather give you a "
         "confirmed timetable than an optimistic one.")
    turn(d, "00:06:54", sp("AV"),
         "In writing, please, and before the end of this month. Because the compensation "
         "plan hangs off that release, and Go-to-Market hangs off the compensation plan.")
    turn(d, "00:07:06", sp("LK"),
         "It does. Territory and quota alignment cannot be finalised until the compensation "
         "design is signed. I am green today but I am green because my dependencies have not "
         "bitten yet, not because I am insulated from them.")
    turn(d, "00:07:22", sp("AV"),
         "That is an honest green. I will take it.")
    turn(d, "00:07:26", sp("SL"),
         "For the record, the committee accepted all seven ratings as submitted. The three "
         "ambers were challenged and a mitigation was requested in each case.")

    # ================================================================ 4
    d.add_heading("00:08 - Value prioritisation framework", level=1)
    turn(d, "00:08:02", sp("AV"),
         f"{C.nm('advisor')}, this is yours.")
    turn(d, "00:08:06", sp("CF"),
         "Thank you. The proposal is that this office adopts a single value prioritisation "
         "framework as the basis for sequencing, and that nothing gets sequenced any other "
         "way. The logic is not complicated. In a programme with more than twenty "
         "workstreams, roughly a fifth of the opportunities carry most of the accretive "
         "value. If you do not decide deliberately which fifth, you end up spreading "
         "attention evenly, which in practice means the loudest workstream wins.")
    turn(d, "00:08:44", sp("CF"),
         "The scoring is through a customer-first lens. Not what is easiest for us "
         "internally, what changes the customer's experience of the combination.")
    turn(d, "00:08:55", sp("KM"),
         "Can I ask a practical question. If my initiative scores low, does that mean it "
         "stops, or does it mean it queues?")
    turn(d, "00:09:04", sp("CF"),
         "It queues. The framework sequences, it does not cancel. But it does mean that when "
         "two things compete for the same person in the same week, the score decides and not "
         "the escalation volume.")
    turn(d, "00:09:18", sp("KM"), "Then I am comfortable.")
    turn(d, "00:09:21", sp("AV"),
         "Anyone against? ... Then that is decision B zero one. The value prioritisation "
         "framework is adopted as the single basis for sequencing integration work, "
         "effective today. It applies to every workstream, no exceptions, and I will not "
         "entertain a sequencing argument that is not made through it.")
    turn(d, "00:09:44", sp("SL"), "Recorded, B zero one.")

    # ================================================================ 5
    d.add_heading("00:10 - Reporting and the single source of truth", level=1)
    turn(d, "00:10:02", sp("SL"),
         "This next one is mine and I am going to be blunt about it. Status is reaching this "
         "office in six different formats. Spreadsheets, slides, Word documents, wiki pages, "
         "mail, and a couple of things that are screenshots of other things. My team spends "
         "most of the week reconciling those against each other, which is time that was "
         "supposed to go into analysis.")
    turn(d, "00:10:33", sp("AV"),
         "How much of the week?")
    turn(d, "00:10:36", sp("SL"),
         "Honestly, most of it. And the reconciliation does not produce anything. It just "
         "gets us back to a single set of numbers that we should have had on Monday.")
    turn(d, "00:10:48", sp("DO"),
         "It also creates a second problem, which is that once there are six versions, the "
         "one that gets quoted is whichever one somebody happened to open. We have already "
         "had a figure quoted from a slide that was two weeks behind the tracker.")
    turn(d, "00:11:05", sp("AV"),
         "So what are you asking for?")
    turn(d, "00:11:08", sp("SL"),
         "That the integration tracker is mandated as the single source of truth for tasks, "
         "milestones and dates. Slides can carry commentary. They cannot carry a number that "
         "contradicts the tracker. If a lead disagrees with the tracker, the fix is to "
         "correct the tracker, not to publish a different figure somewhere else.")
    turn(d, "00:11:32", sp("TB"),
         "I would support that. It also removes an argument I keep having.")
    turn(d, "00:11:38", sp("VO"),
         "One caveat. The tracker has to actually be current, or we will all be quoting "
         "something wrong together instead of separately.")
    turn(d, "00:11:48", sp("SL"),
         "Agreed. Which is on me and on the leads, weekly, Thursday twelve hundred.")
    turn(d, "00:11:55", sp("AV"),
         "Then decision B zero two. The integration tracker is the single source of truth "
         "for tasks, milestones and dates, effective today. Where any other document "
         "disagrees, the tracker governs. Implementation sits with "
         f"{C.nm('imo_pmo')}.")
    turn(d, "00:12:14", sp("SL"), "Recorded, B zero two.")

    # ================================================================ 6
    d.add_heading("00:12 - Cross-workstream dependencies", level=1)
    turn(d, "00:12:22", sp("DO"),
         "Which brings me to the thing I think is actually the biggest risk in this "
         "programme, and it is not on anybody's risk register as a single item because it "
         "does not belong to anybody.")
    turn(d, "00:12:36", sp("AV"), "Go on.")
    turn(d, "00:12:38", sp("DO"),
         "Everybody in this programme understands their own workstream. Nobody has a "
         "complete picture of how their work lands on somebody else's. That is not a "
         "criticism, it is a property of a programme with this many moving parts and "
         "thousands of people across the world. But the consequence is that handovers slip "
         "quietly. The receiving workstream finds out late, and by then the date has already "
         "gone.")
    turn(d, "00:13:09", sp("CF"),
         "We see this pattern consistently. The dependencies between workstreams are where "
         "large integrations lose time, more than inside any single workstream.")
    turn(d, "00:13:20", sp("RW"),
         "It is exactly what I described earlier with D zero one. Tax has a date. I have not "
         "agreed it. Right now there is nothing in the process that forces that "
         "disagreement into the open.")
    turn(d, "00:13:34", sp("DO"),
         "So the proposal is simple. Every cross-workstream handover date in the dependency "
         "register has to be confirmed by both leads. Not the giving side alone. If it is "
         "confirmed by one side only, it is flagged, and it appears in this pack until it is "
         "resolved.")
    turn(d, "00:13:55", sp("ND"),
         "That will produce some uncomfortable conversations.")
    turn(d, "00:13:59", sp("DO"), "That is the point.")
    turn(d, "00:14:02", sp("AV"),
         "Then let it produce them. Decision B zero three. Both workstream leads must "
         "confirm every cross-workstream handover date in the dependency register. "
         "Unconfirmed handovers are reported to this committee until they are resolved. "
         f"{C.nm('imo_mgr')} implements.")
    turn(d, "00:14:22", sp("AV"),
         f"And {C.nm('imo_mgr')}, I want the refreshed dependency map circulated to all "
         "leads. Not tabled here, sent to them, so nobody can say they had not seen it.")
    turn(d, "00:14:34", sp("DO"), "I will get it out next week.")
    turn(d, "00:14:37", sp("SL"), "Recorded, B zero three, and an action on the map.")

    # ================================================================ 7
    d.add_heading("00:15 - Value realisation and synergy validation", level=1)
    turn(d, "00:15:03", sp("AV"),
         f"{C.nm('synergy')}, value realisation.")
    turn(d, "00:15:07", sp("MH"),
         f"The pack in front of you shows secured synergy of USD {C.SYN_SECURED_S01:,} "
         "million run-rate. I need to say something about that number before anybody quotes "
         "it outside this room.")
    turn(d, "00:15:22", sp("AV"), "Please do.")
    turn(d, "00:15:24", sp("MH"),
         f"Of that, Finance can currently stand behind USD {C.SYN_VALIDATED_S01} million. "
         "The rest is secured in the sense that a workstream has an initiative, an owner and "
         "a business case, but it has not been validated against the locked baseline.")
    turn(d, "00:15:44", sp("RW"),
         "And I am not going to validate most of it this month, because the revenue "
         "initiatives in particular do not have a measurement basis yet. I can tell you what "
         "we intend to sell. I cannot tell you what of that would not have been sold anyway.")
    turn(d, "00:16:02", sp("AV"),
         "So the pack overstates it.")
    turn(d, "00:16:05", sp("MH"),
         "The pack reports two different things under one heading. That is my error in "
         "presentation and I would like to fix it properly rather than footnote it.")
    turn(d, "00:16:16", sp("AV"),
         "I would rather have a small number I can defend than a large one I cannot. What "
         "happens the first time somebody outside this programme repeats the larger figure "
         "and Finance will not confirm it?")
    turn(d, "00:16:30", sp("RW"), "Then we have a credibility problem that we cannot undo.")
    turn(d, "00:16:35", sp("AV"),
         "Then decision B zero six. From the next session onward, only Finance-validated "
         "initiatives count towards secured synergy in anything this committee sees or "
         "issues. Everything else is reported separately as pipeline and labelled as such.")
    turn(d, "00:16:55", sp("MH"),
         "That will make the number go down before it goes up.")
    turn(d, "00:16:59", sp("AV"),
         "Good. That is what an honest number does. And reconcile the reported figure to the "
         "validated one before session two, so we know exactly how big the gap is.")
    turn(d, "00:17:11", sp("MH"), "I will have it before the papers go out.")
    turn(d, "00:17:15", sp("SL"), "Recorded, B zero six, and an action on the reconciliation.")

    # ================================================================ 8
    d.add_heading("00:18 - Risks above the escalation threshold", level=1)
    turn(d, "00:18:04", sp("AV"),
         f"{C.nm('risk')}, risks. Only the ones above threshold, please, we do not need the "
         "long list.")
    turn(d, "00:18:12", sp("PN"),
         f"The charter sets the escalation threshold at severity {C.ESCALATION_THRESHOLD}, "
         "likelihood times impact. Four risks are escalated to this committee.")
    esc = [r for r in C.RISKS if r[11] == "Steering Committee"]
    ts = ["00:18:24", "00:19:02", "00:19:41", "00:20:18"]
    for t, r in zip(ts, esc):
        turn(d, t, sp("PN"),
             f"{r[0].replace('R-', 'R zero ')}. {r[1]}. Likelihood {r[5]}, impact {r[6]}, "
             f"severity {C.sev(r)}, band {C.band(C.sev(r)).lower()}. Owner "
             f"{C.nm(r[4])}. Mitigation: {r[7].lower()}. Trend is {r[10].lower()}.")
    turn(d, "00:20:58", sp("AV"),
         "On R zero four, the retention one. What is the actual exposure?")
    turn(d, "00:21:05", sp("VO"),
         "The exposure is that we lose people from the combined field before the coverage "
         "model is in force, and the coverage model is what tells them what their job is. "
         "Uncertainty is the thing that makes them leave, not the deal.")
    turn(d, "00:21:22", sp("LK"),
         "Which argues for bringing the coverage model communication forward rather than "
         "waiting for it to be perfect.")
    turn(d, "00:21:31", sp("AV"), "Note that. On R eleven, the debt one, is that mine or "
                                  "Finance's?")
    turn(d, "00:21:38", sp("RW"),
         "Mine to monitor, yours to decide if it bites. If discretionary integration spend "
         "gets constrained, the answer is re-phasing cost to achieve for things off the "
         "critical path, not cutting them.")
    turn(d, "00:21:54", sp("AV"),
         "Understood. I accept all four mitigations as presented. And I want the threshold "
         "applied without filtering. If something reaches "
         f"{C.ESCALATION_THRESHOLD}, it appears in the pack. I do not want a judgement call "
         "in between the register and this room about what is worth my attention.")
    turn(d, "00:22:14", sp("PN"), "Understood, no filtering.")

    # ================================================================ 9
    d.add_heading("00:23 - Day 100 scope and deferred items", level=1)
    turn(d, "00:23:07", sp("AV"),
         f"Last substantive item. Day 100 is {C.en(C.DAY100)}. That date is a commitment to "
         "the sponsor. I have been asked whether we should be looking at the scope of what "
         "we deliver by then.")
    turn(d, "00:23:24", sp("DO"),
         "My view is that we cannot take that decision today, because the biggest single "
         "input to it is the ERP blueprint recovery plan, and we do not have it.")
    turn(d, "00:23:36", sp("TB"),
         "Correct. Deciding what to move without knowing what the recovery costs would be "
         "guessing.")
    turn(d, "00:23:44", sp("AV"),
         "Agreed. Then the Day 100 scope decision is deferred to session two, and the "
         "condition for taking it is that the recovery plan is in front of us with a dated "
         "forecast. Is there anything else we should be deferring rather than pretending to "
         "decide?")
    turn(d, "00:24:03", sp("DO"),
         f"{C.nm('WS7')} asked me to raise the site consolidation plan for the region. The "
         "lease expiry analysis is not finished, so approving a footprint plan on top of it "
         "would be approving an assumption.")
    turn(d, "00:24:19", sp("AV"),
         "Then that is deferred to session three, and it comes back when the analysis is "
         "complete. Not before.")
    stage(d, "00:24:29", "brief overlapping speech, two participants")
    turn(d, "00:24:34", sp("KM"),
         "[inaudible] ... whether the flagship event date has any flexibility.")
    turn(d, "00:24:41", sp("LK"),
         f"None. {C.en(C.DELL_EMC_WORLD)} is fixed and external. Customers have booked "
         "flights. It is the one date in this programme that will not move for us.")
    turn(d, "00:24:54", sp("AV"),
         "Then treat it as a forcing function rather than a problem. Anything that has to be "
         "true before that event should be planned backwards from it.")

    # ================================================================ 10
    d.add_heading("00:25 - Actions, next meeting and close", level=1)
    turn(d, "00:25:11", sp("AV"),
         f"{C.nm('imo_pmo')}, read the actions back so we all leave with the same list.")
    acts = [a for a in C.ACTIONS if "session 01" in a[3]]
    ts2 = ["00:25:18", "00:25:34", "00:25:49", "00:26:05"]
    for t, a in zip(ts2, acts):
        turn(d, t, sp("SL"),
             f"{a[0].replace('OP-', 'Action O P zero ')}. {a[1]}. Owner "
             f"{C.nm(a[4])}. Due {C.en(a[6])}.")
    turn(d, "00:26:22", sp("AV"),
         f"And to be explicit on the first one, that action sits with {C.nm('imo_mgr')}. "
         "If it needs to move to somebody else, it moves in writing and the register gets "
         "updated, not just a conversation in a corridor.")
    turn(d, "00:26:38", sp("DO"), "Understood.")
    turn(d, "00:26:41", sp("AV"),
         f"Session two is {C.en(C.STEERCO_02)}. Focus topics: the Day 100 scope decision, "
         "the ERP blueprint recovery plan, and the works council consultation timetable. "
         f"Papers to members by {C.en(C.D_WS_ONEPAGER)}, please, not the night before.")
    turn(d, "00:27:01", sp("SL"), "Papers by the twenty-eighth. Noted.")
    turn(d, "00:27:06", sp("AV"),
         "Last thing. Four decisions today, four actions, four risks accepted, two items "
         "deferred with conditions attached to both. That is a working committee. Please "
         "keep it that way, and please do not save things up for the meeting. If something "
         "breaks on a Tuesday, I want to know on the Tuesday.")
    turn(d, "00:27:29", sp("AV"),
         "Thank you all. That is session one closed.")
    stage(d, "00:27:35", "general acknowledgement, meeting closes")
    stage(d, "00:27:52", "recording continues, informal conversation between two "
                         "participants, not part of the meeting")
    stage(d, "01:32:41", "recording ends")

    d.add_paragraph()
    note(d, f"End of transcript. Reviewed against the recording by {C.nm('imo_pmo')} on "
            f"{C.en(C.STEERCO_01)}. The approved minutes of this meeting were signed on "
            f"{C.en(C.STEERCO_01_SIGNED)} and are the binding record; where a reader needs "
            f"the decision as taken rather than the discussion around it, the minutes "
            f"govern.")
    d.add_paragraph()
    note(d, "Synthetic document. The case anchoring, meaning the entities, the transaction "
            "value, the close date, the governance structure and the workstream cut, follows "
            "publicly available sources. All speakers, statements, figures, risks, decisions "
            "and synergy values are invented, internally consistent with the rest of this "
            "corpus, and are not a representation of anything any party actually said or did.")
    note(d, "Sources for the case anchoring: " + "; ".join(C.SOURCES))

    d.save(C.OUT / "DellEMC_VCIO_SteerCo_Transcript_Session01_2016-09-15.docx")


if __name__ == "__main__":
    build()
    print("transcript done")
